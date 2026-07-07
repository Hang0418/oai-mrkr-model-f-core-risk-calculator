#!/usr/bin/env python3
"""Generate a manuscript draft DOCX for the OAI landmark project."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript" / "OAI_24m_landmark_dynamic_prediction_manuscript_draft.docx"
TABLES = ROOT / "results" / "tables"
FIGS = ROOT / "results" / "figures"


def fmt(x, digits=3):
    if pd.isna(x):
        return ""
    if isinstance(x, str):
        return x
    return f"{float(x):.{digits}f}"


def pct(x):
    if pd.isna(x):
        return ""
    return f"{100 * float(x):.1f}%"


def pvalue(x):
    if pd.isna(x):
        return ""
    x = float(x)
    return "<0.001" if x < 0.001 else f"{x:.3f}"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(8)


def add_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, h in zip(hdr, headers):
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, val in zip(cells, row):
            set_cell_text(cell, val)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_figure(doc, path, caption):
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.2))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    model = pd.read_csv(TABLES / "oai_docx_plan_model_comparison.csv")
    coef = pd.read_csv(TABLES / "oai_docx_plan_cox_coefficients.csv")
    group_hr = pd.read_csv(TABLES / "oai_docx_plan_symptom_structure_group_hr.csv")
    risk = pd.read_csv(TABLES / "oai_revision_model_e_clinical_risk_strata.csv")
    inc = pd.read_csv(TABLES / "oai_revision_model_c_vs_e_incremental_value.csv")
    incl = pd.read_csv(TABLES / "oai_revision_included_vs_excluded.csv")
    logit = pd.read_csv(TABLES / "oai_revision_fixed_60m_logistic_sensitivity.csv")
    ph = pd.read_csv(TABLES / "oai_revision_model_e_ph_test.csv")
    cal = pd.read_csv(TABLES / "oai_revision_model_e_calibration_summary.csv")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Dynamic clinical-radiographic prediction of knee replacement risk using a 24-month landmark model in the Osteoarthritis Initiative"
    )
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    sub = doc.add_paragraph("Manuscript draft generated from the current reproducible analysis outputs")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "Background: Static baseline prediction models may miss clinically meaningful changes in symptoms and radiographic structure during follow-up in knee osteoarthritis. We evaluated whether a 24-month landmark dynamic clinical-radiographic model improves prediction of future target-knee knee replacement or total knee arthroplasty (KR/TKA) in the Osteoarthritis Initiative (OAI).",
    )
    add_para(
        doc,
        "Methods: We conducted a knee-level landmark analysis among OAI knees that remained at risk at 24 months. Sequential Cox landmark models were developed from basic demographic predictors to a full dynamic clinical-radiographic model incorporating baseline symptoms, baseline KL/JSN measures, and 0-24 month symptom and structural changes. Robust standard errors clustered by participant accounted for within-person correlation. Performance was assessed using Harrell's C-index, time-dependent AUC at 24, 60, and 96 months after landmark, Brier score, calibration, decision curve analysis, 100 patient-level bootstrap repetitions, and 50 repeated patient-level train/test splits.",
    )
    add_para(
        doc,
        "Results: The common complete-case analysis set included 3,066 knees from 1,640 participants, with 559 post-landmark KR/TKA events. The full dynamic clinical-radiographic model had the best performance, with apparent C-index 0.770, optimism-corrected C-index 0.764, repeated split C-index 0.758, 60-month AUC 0.797, and 60-month Brier score 0.111. Baseline KL grade, medial and lateral JSN worsening, and pain worsening were the most interpretable predictors. Symptom-structure discordance analysis showed a marked risk gradient, with high pain/high structural damage having the highest KR/TKA risk.",
    )
    add_para(
        doc,
        "Conclusions: A 24-month dynamic clinical-radiographic landmark model substantially improved KR/TKA risk prediction over demographic or baseline symptom models. Structural severity and JSN progression were dominant predictors, while pain worsening added complementary symptom information. External validation and flexible modeling for non-proportional hazards are warranted before clinical deployment.",
    )

    add_heading(doc, "Introduction", 1)
    add_para(
        doc,
        "Knee osteoarthritis is a leading cause of pain, functional limitation, and knee replacement. Accurate risk stratification for future KR/TKA may help guide follow-up intensity, shared decision-making, and selection of individuals for disease-modifying or surgical evaluation pathways. Many prediction studies rely on baseline information alone, even though clinical decisions are often made after repeated follow-up visits when both symptoms and radiographic structure have evolved.",
    )
    add_para(
        doc,
        "A landmark design is well suited to this clinical scenario. At a fixed follow-up time, individuals who remain at risk can be re-assessed using information accumulated up to that time, and future event risk can then be estimated from the landmark onward. In knee osteoarthritis, this approach allows baseline disease burden and 0-24 month changes in pain, function, KL grade, and joint-space narrowing to be evaluated together.",
    )
    add_para(
        doc,
        "Symptoms and radiographic structure are also known to be discordant in many patients. Some knees have high pain with limited structural damage, whereas others show advanced radiographic osteoarthritis with relatively low pain. Whether such symptom-structure phenotypes provide useful prognostic information for KR/TKA risk remains clinically important.",
    )
    add_para(
        doc,
        "We therefore used OAI data to develop and internally validate a 24-month landmark dynamic clinical-radiographic model for target-knee KR/TKA. We compared sequential models, assessed calibration and clinical utility, and examined symptom-structure discordance as an explanatory prognostic framework.",
    )

    add_heading(doc, "Methods", 1)
    add_heading(doc, "Data source and study design", 2)
    add_para(
        doc,
        "We used the Osteoarthritis Initiative, a multicenter longitudinal observational cohort of individuals with or at risk for knee osteoarthritis. The present analysis was designed as a knee-level 24-month landmark prediction study. The analysis unit was the knee; participants could contribute one or both knees. All resampling, train/test splitting, and robust variance estimation accounted for participant-level clustering.",
    )
    add_heading(doc, "Landmark population and outcome", 2)
    add_para(
        doc,
        "The landmark was defined at 24 months. Eligible knees had not undergone target-knee KR/TKA before or at the landmark and had positive follow-up time after the landmark. The primary outcome was target-knee KR/TKA after 24 months. Follow-up time was calculated as event or censoring month minus 24 months.",
    )
    add_heading(doc, "Predictors and model sequence", 2)
    add_para(
        doc,
        "Five sequential Cox landmark models were specified. Model A included age, sex, BMI, and knee side. Model B added baseline WOMAC pain, function, and stiffness. Model C added baseline radiographic KL grade and medial/lateral JSN. Model D added 0-24 month symptom changes to the baseline symptom model. Model E, the prespecified main model, added symptom changes and structural changes to the baseline clinical-radiographic model.",
    )
    add_heading(doc, "Validation and performance assessment", 2)
    add_para(
        doc,
        "All Model A-E comparisons were performed on a common complete-case set to avoid differences in sample composition. Cox models used cluster-robust standard errors by participant. We estimated apparent C-index, 100 patient-level bootstrap optimism correction, and 50 repeated patient-level train/test split C-index. Time-dependent AUC and Brier scores were evaluated at 24, 60, and 96 months after landmark. Calibration and decision curve analysis were emphasized at 60 months.",
    )
    add_heading(doc, "Sensitivity and explanatory analyses", 2)
    add_para(
        doc,
        "We assessed the proportional hazards assumption for Model E, performed a fixed-horizon 60-month logistic sensitivity analysis, compared included and excluded eligible knees, tested the incremental value of Model E over Model C, and created clinical risk strata using predicted 60-month risk thresholds of <5%, 5-15%, 15-30%, and >30%. Symptom-structure discordance was assessed by classifying knees into low pain/low structure, high pain/low structure, low pain/high structure, and high pain/high structure groups at the landmark.",
    )

    add_heading(doc, "Results", 1)
    add_heading(doc, "Analysis population", 2)
    add_para(
        doc,
        "The common complete-case analysis set included 3,066 knees from 1,640 participants, with 559 KR/TKA events after the landmark. Compared with excluded eligible knees, included knees had higher event rates and greater baseline clinical and structural burden, indicating that imaging-complete analyses may select a higher-risk subset.",
    )

    add_heading(doc, "Sequential model performance", 2)
    add_para(
        doc,
        "Performance increased progressively as symptoms, baseline imaging, and dynamic change variables were added. Model E performed best across discrimination, time-dependent AUC, Brier score, bootstrap internal validation, and repeated patient-level split validation.",
    )
    rows = []
    for _, r0 in model.iterrows():
        rows.append([
            r0["model"].replace("_", " "),
            fmt(r0["c_index"]),
            fmt(r0["optimism_corrected_c_index"]),
            fmt(r0["split_c_index_mean"]),
            fmt(r0["auc_60m"]),
            fmt(r0["brier_60m"]),
        ])
    add_table(doc, rows, ["Model", "C-index", "Boot-corrected C", "Split C", "AUC 60m", "Brier 60m"])

    add_heading(doc, "Main model predictors", 2)
    add_para(
        doc,
        "In Model E, baseline KL grade and 0-24 month JSN worsening were the strongest structural predictors. Pain worsening also remained independently associated with higher future KR/TKA risk, supporting the value of dynamic symptom information alongside radiographic progression.",
    )
    m_e = coef[coef["model"] == "E_dynamic_clinical_imaging"].copy()
    keep_terms = [
        "pain_change",
        "kl_0",
        "kl_change",
        "jsn_medial_change",
        "jsn_lateral_change",
        "stiffness_0",
        "bmi",
    ]
    m_e = m_e[m_e["term"].isin(keep_terms)]
    rows = []
    for _, r0 in m_e.iterrows():
        rows.append([
            r0["term"],
            fmt(r0["hazard_ratio"]),
            f'{fmt(r0["ci_lower_95"])}-{fmt(r0["ci_upper_95"])}',
            pvalue(r0["p_value"]),
        ])
    add_table(doc, rows, ["Predictor", "HR", "95% CI", "P value"])

    add_heading(doc, "Calibration, risk strata, and clinical utility", 2)
    add_para(
        doc,
        f'Model E showed acceptable 60-month calibration overall. Mean predicted risk was {pct(cal.loc[0, "mean_predicted_risk"])}, and observed Kaplan-Meier risk was {pct(cal.loc[0, "observed_km_risk"])}. Calibration slope was approximately {fmt(cal.loc[0, "calibration_slope_survival"])}. The highest risk decile slightly underestimated observed risk, suggesting possible need for recalibration in the highest-risk group.',
    )
    rows = []
    for _, r0 in risk.iterrows():
        rows.append([
            r0["risk_group"],
            int(r0["knees"]),
            int(r0["participants"]),
            int(r0["events_by_60m"]),
            pct(r0["mean_predicted_60m_risk"]),
            pct(r0["observed_km_60m_risk"]),
        ])
    add_table(doc, rows, ["Risk group", "Knees", "Participants", "Events by 60m", "Predicted", "Observed"])
    add_para(
        doc,
        "Decision curve analysis demonstrated positive net benefit for Model E across 2%-30% risk thresholds, with stronger advantage over treat-all and treat-none strategies at thresholds of 6% and above.",
    )

    add_heading(doc, "Incremental value and sensitivity analyses", 2)
    add_para(
        doc,
        f'Model E significantly improved over Model C on the common complete-case set (likelihood-ratio chi-square {fmt(inc.loc[0, "lr_chisq"])}, df={int(inc.loc[0, "df"])}, P<0.001). C-index increased by {fmt(inc.loc[0, "delta_c_index"])}, fixed-horizon 60-month AUC increased by {fmt(inc.loc[0, "delta_auc_60m_fixed_logistic"])}, and AIC decreased by {fmt(abs(inc.loc[0, "delta_aic"]))}.',
    )
    rows = []
    for _, r0 in logit.iterrows():
        rows.append([
            r0["model"],
            int(r0["n"]),
            int(r0["events_60m"]),
            fmt(r0["auc_60m"]),
            fmt(r0["aic"]),
        ])
    add_table(doc, rows, ["Fixed-horizon model", "N", "Events", "AUC 60m", "AIC"])
    add_para(
        doc,
        "The proportional hazards test for Model E was globally significant, indicating non-proportionality for at least some predictors. Therefore, Cox results should be interpreted as an interpretable landmark model, with fixed-horizon and flexible survival models used as sensitivity analyses before final submission.",
    )

    add_heading(doc, "Symptom-structure discordance", 2)
    add_para(
        doc,
        "Symptom-structure groups showed a strong gradient in KR/TKA risk. Compared with low pain/low structure, high pain/low structure had approximately two-fold higher risk, low pain/high structure had approximately four-fold higher risk, and high pain/high structure had approximately eight-fold higher risk.",
    )
    rows = []
    for _, r0 in group_hr[group_hr["term"].str.contains("symptom_structure_group", regex=False)].iterrows():
        label = r0["term"].replace("symptom_structure_group", "")
        rows.append([
            label,
            fmt(r0["hazard_ratio"]),
            f'{fmt(r0["ci_lower_95"])}-{fmt(r0["ci_upper_95"])}',
            pvalue(r0["p_value"]),
        ])
    add_table(doc, rows, ["Group vs low pain/low structure", "HR", "95% CI", "P value"])

    add_heading(doc, "Figures", 1)
    add_figure(
        doc,
        FIGS / "oai_docx_plan_calibration_60m.png",
        "Figure 1. Calibration of the Model E 60-month KR/TKA risk after the 24-month landmark.",
    )
    add_figure(
        doc,
        FIGS / "oai_docx_plan_decision_curve_60m.png",
        "Figure 2. Decision curve analysis for Model E at 60 months after the landmark.",
    )
    add_figure(
        doc,
        FIGS / "oai_docx_plan_km_by_risk_group.png",
        "Figure 3. KR/TKA-free survival by tertile-based Model E predicted risk group.",
    )

    add_heading(doc, "Discussion", 1)
    add_para(
        doc,
        "In this OAI 24-month landmark analysis, a dynamic clinical-radiographic model incorporating baseline symptoms, baseline radiographic severity, and 0-24 month symptom and structural changes substantially improved prediction of future KR/TKA. The model outperformed demographic, baseline symptom, baseline clinical-radiographic, and dynamic clinical-only models across discrimination, time-dependent AUC, Brier score, bootstrap validation, and repeated patient-level split validation.",
    )
    add_para(
        doc,
        "The most clinically interpretable finding is that structural information dominated prediction, but symptom worsening added complementary prognostic value. Baseline KL grade was the strongest baseline predictor, while medial and lateral JSN worsening captured dynamic structural progression. Pain worsening was independently associated with subsequent KR/TKA, suggesting that evolving symptom burden remains relevant even after accounting for imaging.",
    )
    add_para(
        doc,
        "The symptom-structure discordance analysis provides a clinically intuitive framework. Low pain/high structure carried greater risk than high pain/low structure, indicating that radiographic severity may be more strongly linked to eventual KR/TKA than pain alone. However, the highest risk occurred when high pain and high structural damage coexisted, supporting the use of combined symptom and structure information in follow-up decisions.",
    )
    add_para(
        doc,
        "The findings have potential clinical implications. A 24-month dynamic prediction model could support individualized risk communication after a period of observation, identify knees requiring closer monitoring, and help select candidates for intensified conservative management or surgical counseling. Decision curve analysis suggested positive net benefit across clinically plausible risk thresholds.",
    )
    add_para(
        doc,
        "Several limitations should be considered. First, this is an OAI-only analysis with robust internal validation but no definitive external validation. Second, complete-case imaging requirements reduced the analysis set and selected a higher-risk subset. Third, KR/TKA is influenced by symptoms, structural disease, patient preferences, and healthcare access. Fourth, the proportional hazards assumption was globally violated, so additional fixed-horizon or flexible survival sensitivity analyses should be included in the final submission. Finally, death or other competing risks were not explicitly modeled in the current dataset version.",
    )
    add_para(
        doc,
        "In conclusion, an OAI 24-month dynamic clinical-radiographic landmark model provides robust internally validated prediction of future KR/TKA. Baseline structural severity, JSN worsening, and pain worsening jointly identify high-risk knees. The model and symptom-structure framework warrant external validation and refinement before clinical application.",
    )

    add_heading(doc, "Figure Legends", 1)
    add_para(doc, "Figure 1. Decile-based calibration plot comparing mean predicted and observed 60-month KR/TKA risk for Model E.")
    add_para(doc, "Figure 2. Decision curve analysis comparing Model E with treat-all and treat-none strategies at 60 months after landmark.")
    add_para(doc, "Figure 3. KR/TKA-free survival by predicted-risk group after the 24-month landmark.")

    add_heading(doc, "References", 1)
    refs = [
        "Osteoarthritis Initiative. Publicly available data documentation and study materials.",
        "Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent Reporting of a multivariable prediction model for Individual Prognosis Or Diagnosis (TRIPOD).",
        "Wolbers M, Koller MT, Witteman JCM, Steyerberg EW. Prognostic models with competing risks: methods and application to coronary risk prediction.",
        "Vickers AJ, Elkin EB. Decision curve analysis: a novel method for evaluating prediction models.",
    ]
    for ref in refs:
        add_para(doc, ref)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
