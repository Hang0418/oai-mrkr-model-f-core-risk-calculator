#!/usr/bin/env python3
"""Generate reviewer-guided revised manuscript DOCX."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript" / "OAI_24m_landmark_dynamic_prediction_revised_manuscript.docx"
TABLES = ROOT / "results" / "tables"
FIGS = ROOT / "results" / "figures"


def fmt(x, digits=3):
    if pd.isna(x):
        return ""
    if isinstance(x, str):
        return x
    return f"{float(x):.{digits}f}"


def pct(x, digits=1):
    if pd.isna(x):
        return ""
    return f"{100 * float(x):.{digits}f}%"


def pval(x):
    if pd.isna(x):
        return ""
    x = float(x)
    return "<0.001" if x < 0.001 else f"{x:.3f}"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(tbl.rows[0].cells[i], "F2F4F7")
        set_text(tbl.rows[0].cells[i], h, bold=True, size=8)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_text(cells[i], val, size=8)
    if widths:
        for row in tbl.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()
    return tbl


def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8 if level > 1 else 12)
    p.paragraph_format.space_after = Pt(4)
    return p


def fig(doc, filename, caption):
    path = FIGS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(6.4))
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)


def add_ref_paragraph(doc, refs):
    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{i}. {ref}")
        r.font.size = Pt(9)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)

    perf = pd.read_csv(TABLES / "oai_docx_plan_model_comparison.csv")
    coef = pd.read_csv(TABLES / "oai_docx_plan_cox_coefficients.csv")
    strata = pd.read_csv(TABLES / "oai_revision_model_e_clinical_risk_strata.csv")
    inc = pd.read_csv(TABLES / "oai_revision_model_c_vs_e_incremental_value.csv")
    incl = pd.read_csv(TABLES / "oai_revision_included_vs_excluded.csv")
    logit = pd.read_csv(TABLES / "oai_revision_fixed_60m_logistic_sensitivity.csv")
    ph = pd.read_csv(TABLES / "oai_revision_model_e_ph_test.csv")
    cal = pd.read_csv(TABLES / "oai_revision_model_e_calibration_summary.csv")
    discord = pd.read_csv(TABLES / "oai_docx_plan_symptom_structure_group_hr.csv")
    discord_sum = pd.read_csv(TABLES / "oai_docx_plan_symptom_structure_group_summary.csv")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Dynamic clinical-radiographic prediction of knee replacement risk using a 24-month landmark model in the Osteoarthritis Initiative")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    sub = doc.add_paragraph("Revised manuscript draft following SCI reviewer-oriented revision plan")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    heading(doc, "Abstract", 1)
    para(doc, "Background: Prediction models for knee replacement or total knee arthroplasty (KR/TKA) in knee osteoarthritis often rely on static baseline information. Whether changes in symptoms and radiographic structure observed during follow-up improve future KR/TKA risk prediction remains clinically relevant.")
    para(doc, "Methods: We performed a knee-level 24-month landmark analysis in the Osteoarthritis Initiative. Knees at risk at the landmark were used to develop sequential Cox landmark models from demographic predictors to a full dynamic clinical-radiographic model incorporating baseline WOMAC symptoms, baseline Kellgren-Lawrence (KL) grade and joint-space narrowing (JSN), and 0-24 month changes in symptoms and structure. Robust standard errors were clustered by participant. Model performance was evaluated by Harrell's C-index, time-dependent AUC, Brier score, calibration, decision curve analysis, 100 patient-level bootstrap repetitions, and 50 repeated patient-level train/test splits.")
    para(doc, "Results: The common complete-case analysis set included 3,066 knees from 1,640 participants, with 559 post-landmark KR/TKA events. The full dynamic clinical-radiographic model achieved the best performance (C-index 0.770; optimism-corrected C-index 0.764; repeated split C-index 0.758; 60-month AUC 0.797; 60-month Brier score 0.111). Compared with the baseline clinical-radiographic model, the dynamic model improved C-index by 0.054 and fixed-horizon 60-month AUC by 0.060. Baseline KL grade, medial and lateral JSN worsening, and pain worsening were the main interpretable predictors. Symptom-structure groups showed a graded KR/TKA risk pattern, with high pain/high structural damage having the highest risk.")
    para(doc, "Conclusions: A 24-month dynamic clinical-radiographic landmark model improved internally validated KR/TKA prediction in OAI. Structural severity and JSN progression were dominant predictors, while pain worsening contributed complementary prognostic information. External validation and flexible modeling for non-proportional hazards are needed before clinical use.")

    heading(doc, "Introduction", 1)
    para(doc, "Knee osteoarthritis is a heterogeneous condition in which pain, functional limitation, and structural progression influence treatment decisions, including KR/TKA. Accurate risk stratification could help identify knees requiring closer monitoring, guide shared decision-making, and support selection of individuals for prevention or surgical referral pathways.")
    para(doc, "Recent models have shown that demographic, clinical, radiographic, and machine-learning features can predict total knee replacement risk in multicentre osteoarthritis cohorts [1,2]. Deep-learning analyses of MRI also support the prognostic value of imaging information [3]. However, many prediction studies remain anchored to baseline characteristics, even though clinical decisions are commonly made after repeated follow-up visits.")
    para(doc, "A landmark prediction design addresses this gap by re-estimating risk at a clinically meaningful follow-up time using information accrued before that time. In knee osteoarthritis, a 24-month landmark allows early symptom trajectories and radiographic changes to be combined with baseline disease burden. This design is especially relevant because pain and imaging findings are frequently discordant [6,7].")
    para(doc, "We developed and internally validated a 24-month dynamic clinical-radiographic landmark model for future target-knee KR/TKA in OAI. We compared sequential models, quantified the incremental value of dynamic symptom and structural changes, evaluated calibration and decision utility, and examined symptom-structure discordance as an explanatory clinical phenotype.")

    heading(doc, "Methods", 1)
    heading(doc, "Study design and data source", 2)
    para(doc, "This was a knee-level prediction-model study using OAI data. Each knee was treated as an observation, while within-participant correlation was handled by cluster-robust standard errors and participant-level resampling. The study was reported with reference to TRIPOD+AI principles for prediction-model reporting [4].")
    heading(doc, "Landmark population and outcome", 2)
    para(doc, "The landmark was set at 24 months. Knees that had undergone target-knee KR/TKA before or at the landmark were excluded. Follow-up began at 24 months, and the primary outcome was target-knee KR/TKA after the landmark. Survival time was defined as KR/TKA month minus 24 months for events and last follow-up month minus 24 months for censored knees.")
    heading(doc, "Predictors and sequential models", 2)
    para(doc, "We specified five sequential models. Model A included age, sex, BMI, and knee side. Model B added baseline WOMAC pain, function, and stiffness. Model C added baseline KL grade and medial/lateral JSN. Model D added 0-24 month symptom changes to the baseline symptom model. Model E, the primary model, combined baseline symptoms, baseline radiographic severity, 0-24 month symptom changes, and 0-24 month structural changes.")
    heading(doc, "Missing data and complete-case analysis", 2)
    para(doc, "The primary Model A-E comparison used a common complete-case set to ensure that differences in performance reflected model content rather than changes in sample composition. Model-specific sample sizes were summarized separately. Included and excluded eligible knees were compared to evaluate potential complete-case selection bias.")
    heading(doc, "Statistical analysis", 2)
    para(doc, "Cox landmark models used robust standard errors clustered by participant. Discrimination was assessed using Harrell's C-index and time-dependent AUC at 24, 60, and 96 months after the landmark. Prediction error was evaluated using Brier scores. Internal validation used 100 patient-level bootstrap repetitions and 50 repeated patient-level train/test splits. Calibration and decision curve analysis were evaluated at 60 months.")
    para(doc, "We tested the proportional hazards assumption for Model E using Schoenfeld residuals. Because non-proportionality can affect interpretation of Cox hazard ratios, we also performed a fixed-horizon 60-month logistic sensitivity analysis. Incremental value of Model E over Model C was assessed using a likelihood-ratio test, change in C-index, change in fixed-horizon AUC, and change in AIC.")
    heading(doc, "Symptom-structure discordance", 2)
    para(doc, "At the landmark visit, knees were classified into low pain/low structure, high pain/low structure, low pain/high structure, or high pain/high structure groups. High pain was defined by the upper quartile of landmark WOMAC pain, and high structural damage was defined by advanced radiographic burden (KL grade at least 3 or high JSN burden). Associations with subsequent KR/TKA were estimated using Cox models adjusted for age, sex, BMI, and knee side.")

    heading(doc, "Results", 1)
    fig(doc, "figure1_study_flow_landmark.png", "Figure 1. Study flow and 24-month landmark design. Panel A shows knee-level eligibility and the common complete-case analysis set. Panel B shows how baseline, interim, and 24-month information were used to predict future target-knee KR/TKA.")

    heading(doc, "Analysis population and complete-case selection", 2)
    para(doc, "The common complete-case analysis set included 3,066 knees from 1,640 participants, with 559 post-landmark KR/TKA events. The excluded eligible set included 5,248 knees from 2,717 participants and 272 events. Included knees had higher baseline symptom and structural burden and a higher event rate, suggesting that imaging-complete analyses selected a higher-risk subset.")
    rows = []
    for _, r0 in incl.iterrows():
        group_label = "Included" if str(r0["group"]).startswith("included") else "Excluded"
        rows.append([group_label, int(r0["knees"]), int(r0["participants"]), pct(r0["event_rate"]), fmt(r0["age_mean"], 1), pct(r0["female_pct"]), fmt(r0["bmi_mean"], 1), fmt(r0["pain0_mean"], 2), fmt(r0["kl0_mean"], 2)])
    table(doc, ["Group", "Knees", "Participants", "Event rate", "Age", "Female", "BMI", "Pain 0m", "KL 0m"], rows)

    heading(doc, "Sequential model performance", 2)
    para(doc, "Model performance improved in a stepwise pattern from demographic predictors to dynamic clinical-radiographic prediction. Model E had the best apparent C-index, optimism-corrected C-index, repeated split C-index, 60-month AUC, and 60-month Brier score.")
    rows = []
    for _, r0 in perf.iterrows():
        rows.append([r0["model"].replace("_", " "), fmt(r0["c_index"]), fmt(r0["optimism_corrected_c_index"]), fmt(r0["split_c_index_mean"]), fmt(r0["auc_60m"]), fmt(r0["brier_60m"])])
    table(doc, ["Model", "C-index", "Boot-corrected C", "Split C", "AUC 60m", "Brier 60m"], rows)
    fig(doc, "figure2_model_performance.png", "Figure 2. Incremental model performance across Models A-E. All model comparisons used the same common complete-case analysis set.")

    heading(doc, "Primary Model E predictors and incremental value", 2)
    para(doc, "The primary dynamic clinical-radiographic model identified baseline KL grade, medial and lateral JSN worsening, and pain worsening as the main interpretable predictors. Model E significantly improved over Model C on the common complete-case set.")
    e = coef[coef["model"] == "E_dynamic_clinical_imaging"]
    terms = ["kl_0", "jsn_medial_change", "jsn_lateral_change", "pain_change", "stiffness_0", "bmi"]
    rows = []
    for _, r0 in e[e["term"].isin(terms)].iterrows():
        rows.append([r0["term"], fmt(r0["hazard_ratio"]), f'{fmt(r0["ci_lower_95"])}-{fmt(r0["ci_upper_95"])}', pval(r0["p_value"])])
    table(doc, ["Predictor", "HR", "95% CI", "P value"], rows)
    para(doc, f'Model E improved over Model C with likelihood-ratio chi-square {fmt(inc.loc[0, "lr_chisq"])}, df={int(inc.loc[0, "df"])}, P<0.001. C-index increased by {fmt(inc.loc[0, "delta_c_index"])}, fixed-horizon 60-month AUC increased by {fmt(inc.loc[0, "delta_auc_60m_fixed_logistic"])}, and AIC decreased by {fmt(abs(inc.loc[0, "delta_aic"]))}.')

    heading(doc, "Calibration, decision utility, and clinical risk strata", 2)
    para(doc, f'Model E showed acceptable overall calibration at 60 months. Mean predicted risk was {pct(cal.loc[0, "mean_predicted_risk"])}, observed Kaplan-Meier risk was {pct(cal.loc[0, "observed_km_risk"])}, and the survival calibration slope was {fmt(cal.loc[0, "calibration_slope_survival"])}. Decision curve analysis showed positive net benefit from 2% to 30% thresholds, particularly above 6%.')
    fig(doc, "figure3_calibration_dca.png", "Figure 3. Calibration and decision-curve analysis for Model E at 60 months after the landmark. Calibration labels show risk-decile sample sizes.")
    rows = []
    for _, r0 in strata.iterrows():
        rows.append([r0["risk_group"], int(r0["knees"]), int(r0["participants"]), int(r0["events_by_60m"]), pct(r0["mean_predicted_60m_risk"]), pct(r0["observed_km_60m_risk"])])
    table(doc, ["Risk group", "Knees", "Participants", "Events by 60m", "Predicted risk", "Observed risk"], rows)
    fig(doc, "figure4_clinical_risk_cumulative_incidence.png", "Figure 4. Cumulative KR/TKA incidence by Model E clinical risk strata, truncated at 96 months after the landmark with number-at-risk table.")

    heading(doc, "Sensitivity analyses", 2)
    global_p = pval(ph.loc[ph["term"] == "GLOBAL", "p"].iloc[0])
    global_p_text = "P<0.001" if global_p.startswith("<") else f"P={global_p}"
    para(doc, f'The proportional hazards test for Model E was globally significant (global {global_p_text}), indicating non-proportional hazards for at least some predictors. Fixed-horizon 60-month logistic sensitivity analysis supported the primary finding: Model E achieved AUC {fmt(logit.loc[logit["model"].str.contains("Model E"), "auc_60m"].iloc[0])}, compared with {fmt(logit.loc[logit["model"].str.contains("Model C"), "auc_60m"].iloc[0])} for Model C.')
    para(doc, "Death was not explicitly modeled as a competing event in the current analysis. This limitation is addressed in the Discussion and should be revisited if death or competing-risk data are available in the final analytic dataset.")

    doc.add_page_break()
    heading(doc, "Symptom-structure discordance", 2)
    para(doc, "Symptom-structure groups showed a strong KR/TKA risk gradient. Compared with low pain/low structure, high pain/low structure, low pain/high structure, and high pain/high structure were each associated with higher risk, with the highest risk observed when high pain and high structural damage coexisted.")
    rows = []
    for _, r0 in discord[discord["term"].str.contains("symptom_structure_group", regex=False)].iterrows():
        rows.append([r0["term"].replace("symptom_structure_group", ""), fmt(r0["hazard_ratio"]), f'{fmt(r0["ci_lower_95"])}-{fmt(r0["ci_upper_95"])}', pval(r0["p_value"])])
    table(doc, ["Group vs low pain/low structure", "HR", "95% CI", "P value"], rows)
    fig(doc, "figure5_symptom_structure_discordance.png", "Figure 5. Symptom-structure discordance. Panel A shows observed KR/TKA event rates. Panel B shows adjusted hazard ratios relative to low pain/low structure.")

    heading(doc, "Discussion", 1)
    para(doc, "This 24-month landmark analysis showed that dynamic clinical-radiographic information improved prediction of future target-knee KR/TKA in OAI. The full Model E outperformed demographic, baseline symptom, baseline clinical-radiographic, and dynamic clinical-only models across discrimination, time-dependent AUC, Brier score, and internal validation.")
    para(doc, "These findings extend previous KR/TKA prediction work using multicentre cohorts and machine-learning approaches [1,2]. Unlike static baseline models, the present design used a clinically interpretable follow-up decision point and incorporated change in symptoms and radiographic structure during the first two years. This supports a dynamic risk reassessment strategy rather than a one-time baseline prediction approach.")
    para(doc, "The dominant contribution of KL grade and JSN worsening is consistent with literature emphasizing structural findings and imaging-derived features as important predictors of knee replacement or osteoarthritis progression [1,3,5,10]. Pain worsening nevertheless added complementary information, aligning with evidence that pain and structural damage are related but not interchangeable dimensions of knee osteoarthritis [6,7].")
    para(doc, "The symptom-structure discordance analysis is clinically useful. Low pain/high structure carried greater risk than high pain/low structure, suggesting that structural severity may be more closely linked to eventual KR/TKA than pain alone. However, high pain/high structure identified the highest-risk group, supporting combined symptom and structure assessment during follow-up.")
    para(doc, "Strengths include a prespecified landmark framework, knee-level side-specific analysis, sequential model comparison on a common complete-case set, participant-level bootstrap validation, repeated participant-level train/test splitting, calibration, decision curve analysis, and explicit evaluation of symptom-structure discordance. These features address several reporting and validation concerns raised in recent prediction-model guidance [4,8,9].")
    para(doc, "Several limitations remain. First, the study used OAI only and therefore provides robust internal validation rather than definitive external validation. Second, complete-case imaging requirements selected a higher-risk subset; multiple-imputation or weighting analyses should be considered before submission. Third, proportional hazards assumptions were globally violated, so Cox hazard ratios should be interpreted as an interpretable landmark summary and supported by fixed-horizon or flexible survival sensitivity analyses. Fourth, death was not explicitly modeled as a competing event. Finally, KR/TKA reflects structural disease, symptoms, preferences, surgeon decision-making, and health-system factors.")
    para(doc, "In conclusion, a 24-month dynamic clinical-radiographic landmark model improved internally validated prediction of future KR/TKA in OAI. Baseline structural severity, JSN progression, and pain worsening jointly identified high-risk knees. The model should undergo external validation, recalibration, and competing-risk assessment before clinical deployment.")

    heading(doc, "Data availability", 1)
    para(doc, "OAI data are available through the official OAI/NIMH Data Archive access process subject to data-use requirements. Derived analytic code and variable mappings can be shared where permitted by the OAI data-use agreement.")
    heading(doc, "Code availability", 1)
    para(doc, "The analysis scripts used to generate this draft are stored in the project scripts directory, including scripts for landmark model comparison, sensitivity analyses, figure generation, and manuscript generation.")
    heading(doc, "Ethics statement", 1)
    para(doc, "Ethics approval and informed consent for the original OAI cohort should be cited according to OAI documentation. The present secondary analysis used de-identified data; local ethics requirements should be confirmed before submission.")
    heading(doc, "Competing interests", 1)
    para(doc, "The authors declare no competing interests. [Placeholder to be confirmed before submission.]")

    heading(doc, "References", 1)
    refs = [
        "Liu et al. Prediction models for the risk of total knee replacement: development and validation using data from multicentre cohort studies. The Lancet Rheumatology. 2022.",
        "Mahmoud et al. Predicting total knee replacement at 2 and 5 years in osteoarthritis patients using machine learning. BMJ Surgery, Interventions, & Health Technologies. 2023.",
        "Rajamohan et al. Prediction of total knee replacement using deep learning analysis of knee MRI. Scientific Reports. 2023.",
        "Collins et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024.",
        "Li et al. Predictive models of radiographic progression and pain progression in patients with knee osteoarthritis: data from the FNIH OA biomarkers consortium project. Arthritis Research & Therapy. 2024.",
        "Zhao et al. Identifying significant structural factors associated with knee pain severity in patients with osteoarthritis using machine learning. Scientific Reports. 2024.",
        "Hill et al. The discordance between pain and imaging in knee osteoarthritis. Journal of the American Academy of Orthopaedic Surgeons. 2025.",
        "Joseph et al. Machine learning models for clinical and structural knee osteoarthritis outcomes: review. 2025.",
        "Predictive value of machine learning in knee osteoarthritis progression: systematic review and meta-analysis. Journal of Medical Internet Research. 2025.",
        "Guo et al. Predicting joint space changes in knee osteoarthritis over 6 years: a combined model of TransUNet and XGBoost. Quantitative Imaging in Medicine and Surgery. 2025.",
        "Martel-Pelletier et al. Next-level prediction of structural progression in knee osteoarthritis. 2025.",
    ]
    add_ref_paragraph(doc, refs)

    heading(doc, "Figure legends", 1)
    para(doc, "Figure 1. Study flow and 24-month landmark design. KR/TKA denotes knee replacement or total knee arthroplasty; KL denotes Kellgren-Lawrence grade; JSN denotes joint-space narrowing.")
    para(doc, "Figure 2. Incremental performance across sequential landmark models. All comparisons use the common complete-case analysis set of 3,066 knees.")
    para(doc, "Figure 3. Calibration and decision curve analysis for Model E at 60 months after the landmark.")
    para(doc, "Figure 4. Cumulative KR/TKA incidence by clinical risk strata derived from Model E predicted 60-month risk.")
    para(doc, "Figure 5. Symptom-structure discordance and future KR/TKA risk.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
