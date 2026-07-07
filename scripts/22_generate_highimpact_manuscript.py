from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "results" / "tables"
FIGURES = ROOT / "results" / "figures"
OUT = ROOT / "drafts"
DOCX_OUT = OUT / "OAI_MRKR_high_impact_manuscript_reviewer_strengthened_draft.docx"
LIT_OUT = OUT / "OAI_MRKR_literature_and_analysis_synthesis.md"


def pct(x: float, digits: int = 1) -> str:
    return f"{100 * float(x):.{digits}f}%"


def num(x: float, digits: int = 3) -> str:
    return f"{float(x):.{digits}f}"


def pval(x: float) -> str:
    x = float(x)
    if x < 0.001:
        return f"{x:.1e}"
    return f"{x:.3f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9E2EC")
        borders.append(tag)
    tbl_pr.append(borders)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_table(doc: Document, data: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(data[0]):
        cell = hdr.cells[i]
        set_cell_shading(cell, "EAF1F8")
        set_cell_text(cell, value, bold=True, size=8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[i])
    for row_values in data[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            set_cell_text(cells[i], value, size=8)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
    set_table_borders(table)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 53, 85)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def load_tables():
    model_cmp = pd.read_csv(TABLES / "oai_docx_plan_model_comparison.csv")
    coeff_e = pd.read_csv(TABLES / "oai_docx_plan_cox_coefficients.csv")
    inc = pd.read_csv(TABLES / "oai_revision_model_c_vs_e_incremental_value.csv")
    time_tbl = pd.read_csv(TABLES / "oai_mrkr_plan_time_structure_table.csv")
    fcoef = pd.read_csv(TABLES / "oai_mrkr_plan_model_f_core_coefficients.csv")
    transport = pd.read_csv(TABLES / "oai_mrkr_plan_transport_metrics_by_horizon.csv")
    strata = pd.read_csv(TABLES / "oai_mrkr_plan_mrkr_risk_strata_24m.csv")
    strict = pd.read_csv(TABLES / "oai_mrkr_plan_mrkr_strict_sensitivity_24m.csv")
    outcome = pd.read_csv(TABLES / "oai_mrkr_plan_mrkr_outcome_sensitivity_24m.csv")
    subgroups = pd.read_csv(TABLES / "oai_mrkr_highimpact_mrkr_subgroup_performance_24m.csv")
    cpt = pd.read_csv(TABLES / "mrkr_cpt_laterality_audit_summary.csv")
    table1 = pd.read_csv(TABLES / "oai_mrkr_reviewer_table1_baseline_smd.csv")
    oai_flow = pd.read_csv(TABLES / "oai_reviewer_inclusion_exclusion_flow.csv")
    mrkr_flow = pd.read_csv(TABLES / "mrkr_reviewer_inclusion_exclusion_flow.csv")
    perf_ci = pd.read_csv(TABLES / "oai_mrkr_reviewer_discrimination_ci.csv")
    f_formula = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_f_formula_coefficients.csv")
    f_basehaz = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_f_baseline_hazard.csv")
    f_std = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_f_standardization.csv")
    recal_formula = pd.read_csv(TABLES / "oai_mrkr_reviewer_recalibration_formula.csv")
    ph = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_e_schoenfeld_ph.csv")
    ph_sens = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_e_ph_sensitivity_performance.csv")
    return {
        "model_cmp": model_cmp,
        "coeff_e": coeff_e,
        "inc": inc,
        "time": time_tbl,
        "fcoef": fcoef,
        "transport": transport,
        "strata": strata,
        "strict": strict,
        "outcome": outcome,
        "subgroups": subgroups,
        "cpt": cpt,
        "table1": table1,
        "oai_flow": oai_flow,
        "mrkr_flow": mrkr_flow,
        "perf_ci": perf_ci,
        "f_formula": f_formula,
        "f_basehaz": f_basehaz,
        "f_std": f_std,
        "recal_formula": recal_formula,
        "ph": ph,
        "ph_sens": ph_sens,
    }


def make_literature_synthesis(t) -> None:
    text = f"""# OAI-MRKR literature and analysis synthesis

## Core positioning

This project should be framed as a two-level prediction-model study: (1) a complete OAI 24-month landmark model that uses rich OAI symptom and radiographic change information, and (2) a transport validation and recalibration exercise in MRKR using a deliberately restricted common-variable model. The strongest claim is not that MRKR exactly reproduces OAI, but that the OAI-derived rank-order signal persists in a more heterogeneous real-world public imaging/EHR resource after accounting for baseline risk differences.

## What the recent public-data modeling literature teaches us

- Multimodal OA prediction papers using public cohorts such as OAI commonly combine clinical variables with imaging phenotypes, then report discrimination over clinically meaningful horizons. Tiulpin et al. used radiographs plus clinical covariates to predict OA progression and showed independent-test AUC around 0.79, reinforcing the value of combining imaging and clinical context rather than reporting imaging alone.
- MRI deep-learning work in OAI, including Tolpadi et al. and later baseline-MRI TKR prediction studies, generally emphasizes long-horizon prediction, patient-level splitting, and the gap between high discrimination and clinical usefulness. This supports keeping calibration, risk strata, and decision-curve analyses in the main manuscript.
- Recent OA pain-modeling work using OAI highlights that structural features and symptoms are related but not interchangeable. This supports the current paper's central Model E logic: baseline pain and 0-24 month pain change carry information beyond radiographic severity, while structural change remains a strong surgical-risk signal.
- MRKR is useful because it is a large public real-world radiograph/EHR dataset with pain scores, CPT/ICD data, laterality, KL inference, and arthroplasty hardware labels. Its scale and diversity make it a credible transport-validation setting, but the MRKR outcomes and baseline ascertainment are not identical to OAI adjudicated follow-up.
- Modern prediction-model reporting guidance, including TRIPOD+AI and PROBAST/PROBAST-AI thinking, rewards transparent data provenance, patient-level splitting, calibration assessment, decision-curve analysis, and clear separation between model development, validation, recalibration, and model updating.

## How the analysis was strengthened

- The manuscript now preserves Model E as the primary OAI model: baseline symptoms, baseline KL/JSN, 0-24 month symptom changes, and 0-24 month structural changes.
- MRKR validation uses Model F-core, restricted to variables approximately mappable in both OAI and MRKR: age, sex, knee side, standardized landmark pain, baseline KL, and KL worsening.
- MRKR arthroplasty hardware is retained as the primary outcome because it is side-specific. CPT-based outcomes are reported as sensitivity analyses because many CPT arthroplasty rows are patient-level or lack reliable side-specific modifiers.
- Early MRKR events are handled with strict sensitivity analyses excluding events within 3, 6, and 12 months after the landmark.
- Absolute risk transport is explicitly recalibrated in MRKR using slope plus baseline recalibration. Offset-only recalibration is reported as a diagnostic contrast, not as the preferred approach.

## Key results to carry into the paper

- OAI Model E common-complete analysis: 3,066 knees, 1,640 participants, 559 KR/TKA events; apparent C-index {num(t['model_cmp'].query("model == 'E_dynamic_clinical_imaging'").iloc[0]['c_index'])}; optimism-corrected C-index {num(t['model_cmp'].query("model == 'E_dynamic_clinical_imaging'").iloc[0]['optimism_corrected_c_index'])}; repeated patient-split C-index {num(t['model_cmp'].query("model == 'E_dynamic_clinical_imaging'").iloc[0]['split_c_index_mean'])}; 60-month AUC {num(t['model_cmp'].query("model == 'E_dynamic_clinical_imaging'").iloc[0]['auc_60m'])}.
- Incremental value of Model E over baseline symptoms plus imaging (Model C): likelihood-ratio chi-square {num(t['inc'].iloc[0]['lr_chisq'], 1)}, p={pval(t['inc'].iloc[0]['p_value'])}, delta C-index {num(t['inc'].iloc[0]['delta_c_index'])}, and delta fixed-horizon 60-month AUC {num(t['inc'].iloc[0]['delta_auc_60m_fixed_logistic'])}.
- OAI-MRKR Model F-core OAI training set: 3,104 knees, 1,656 participants, 566 events. MRKR validation set: 3,412 knees, 2,179 patients, 1,140 side-specific hardware events.
- MRKR 24-month validation: original OAI-derived Model F-core C-index {num(t['transport'].query("cohort == 'MRKR' and model == 'Original OAI-derived Model F-core' and horizon_months == 24").iloc[0]['c_index'])}, AUC {num(t['transport'].query("cohort == 'MRKR' and model == 'Original OAI-derived Model F-core' and horizon_months == 24").iloc[0]['auc_horizon'])}; mean predicted risk {pct(t['transport'].query("cohort == 'MRKR' and model == 'Original OAI-derived Model F-core' and horizon_months == 24").iloc[0]['mean_predicted_risk'])} versus observed risk {pct(t['transport'].query("cohort == 'MRKR' and model == 'Original OAI-derived Model F-core' and horizon_months == 24").iloc[0]['observed_km_risk'])}.
- MRKR slope plus baseline recalibration aligned 24-month mean predicted risk {pct(t['transport'].query("cohort == 'MRKR' and model == 'MRKR slope+baseline recalibrated Model F-core' and horizon_months == 24").iloc[0]['mean_predicted_risk'])} with observed risk {pct(t['transport'].query("cohort == 'MRKR' and model == 'MRKR slope+baseline recalibrated Model F-core' and horizon_months == 24").iloc[0]['observed_km_risk'])}, while preserving discrimination.

## References and source pages reviewed

1. MRKR dataset preprint, arXiv: https://arxiv.org/abs/2411.00866
2. Tiulpin et al., multimodal machine learning for knee OA progression, arXiv: https://arxiv.org/abs/1904.06236
3. Tolpadi et al., knee replacement prediction from MRI using deep learning, Scientific Reports: https://www.nature.com/articles/s41598-020-63395-9
4. Rajamohan et al., baseline MRI prediction of future knee replacement, Scientific Reports: https://www.nature.com/articles/s41598-023-33934-1
5. Zhao et al., OAI MRI/semiquantitative feature prediction of knee pain severity, Scientific Reports: https://www.nature.com/articles/s41598-024-65613-0
6. Van Calster et al., calibration in clinical prediction models, BMC Medicine: https://link.springer.com/article/10.1186/s12916-019-1466-7
7. Vickers et al., decision curve analysis interpretation, Diagnostic and Prognostic Research: https://link.springer.com/article/10.1186/s41512-019-0064-7
8. TRIPOD+AI reporting guidance, BMJ: https://www.bmj.com/content/385/bmj-2023-078378

"""
    OUT.mkdir(exist_ok=True)
    LIT_OUT.write_text(text, encoding="utf-8")


def build_docx(t) -> None:
    OUT.mkdir(exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Dynamic clinical-radiographic prediction of knee replacement risk with real-world transport validation in MRKR"
    )
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(20, 45, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run("Reviewer-strengthened manuscript draft | OAI development and MRKR transport validation")
    rr.italic = True
    rr.font.name = "Arial"
    rr.font.size = Pt(9)

    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "Background: Risk prediction for knee replacement in osteoarthritis should reflect both baseline disease burden and dynamic changes in symptoms and joint structure. Most models are developed in deeply phenotyped cohorts, but their transportability to real-world imaging/EHR resources remains uncertain.",
    )
    add_para(
        doc,
        "Methods: We developed 24-month landmark Cox models in the Osteoarthritis Initiative (OAI). The primary OAI model, Model E, included age, sex, body mass index, knee side, baseline WOMAC pain/function/stiffness, baseline KL grade and medial/lateral joint-space narrowing, and 0-24 month changes in symptoms and structure. To support external transport validation, we separately fit Model F-core using only variables approximately mappable in both OAI and the Multicenter Osteoarthritis Study Radiographic Knee Replacement dataset (MRKR): age, sex, knee side, standardized landmark pain, baseline KL grade, and KL worsening. MRKR validation used side-specific arthroplasty hardware as the primary outcome, with CPT-based sensitivity analyses. Performance was assessed using discrimination, time-dependent AUC, calibration, Brier score, decision curves, risk strata, recalibration, early-event exclusions, and subgroup analyses.",
    )
    e = t["model_cmp"].query("model == 'E_dynamic_clinical_imaging'").iloc[0]
    inc = t["inc"].iloc[0]
    mrkr24 = t["transport"].query(
        "cohort == 'MRKR' and model == 'Original OAI-derived Model F-core' and horizon_months == 24"
    ).iloc[0]
    mrkr24_recal = t["transport"].query(
        "cohort == 'MRKR' and model == 'MRKR slope+baseline recalibrated Model F-core' and horizon_months == 24"
    ).iloc[0]
    ci_e = f"{num(e['c_index'])}"
    ci_oai_f = t["perf_ci"].query("cohort_model == 'OAI Model F-core apparent' and metric == 'c_index'").iloc[0]
    ci_mrkr_f = t["perf_ci"].query("cohort_model == 'MRKR original OAI-derived Model F-core' and metric == 'c_index'").iloc[0]
    auc_mrkr_24_ci = t["perf_ci"].query("cohort_model == 'MRKR original OAI-derived Model F-core' and metric == 'auc' and horizon_months == 24").iloc[0]
    add_para(
        doc,
        f"Results: The OAI Model E complete analysis included 3,066 knees from 1,640 participants with 559 knee replacement events. Model E achieved an apparent C-index of {ci_e}, optimism-corrected C-index of {num(e['optimism_corrected_c_index'])}, repeated patient-split C-index of {num(e['split_c_index_mean'])}, and 60-month AUC of {num(e['auc_60m'])}. Compared with baseline symptoms plus imaging, adding 0-24 month symptom and structural change improved fit (likelihood-ratio chi-square {num(inc['lr_chisq'], 1)}, p={pval(inc['p_value'])}) and discrimination (delta C-index {num(inc['delta_c_index'])}; delta 60-month AUC {num(inc['delta_auc_60m_fixed_logistic'])}). Model F-core was trained in 3,104 OAI knees and validated in 3,412 MRKR knees. In MRKR, the original OAI-derived score retained moderate ranking performance at 24 months (C-index {num(mrkr24['c_index'])}, bootstrap 95% CI {num(ci_mrkr_f['ci_lower_95'])}-{num(ci_mrkr_f['ci_upper_95'])}; AUC {num(mrkr24['auc_horizon'])}, 95% CI {num(auc_mrkr_24_ci['ci_lower_95'])}-{num(auc_mrkr_24_ci['ci_upper_95'])}) but underestimated absolute risk (mean predicted {pct(mrkr24['mean_predicted_risk'])}; observed {pct(mrkr24['observed_km_risk'])}). Slope plus baseline recalibration aligned mean predicted risk ({pct(mrkr24_recal['mean_predicted_risk'])}) with observed risk ({pct(mrkr24_recal['observed_km_risk'])}) while preserving discrimination.",
    )
    add_para(
        doc,
        "Conclusions: A dynamic landmark model combining symptoms, radiographic severity, and 0-24 month change improves knee replacement prediction in OAI. When restricted to commonly mappable variables, the OAI-derived signal transports to MRKR for risk ranking, but real-world deployment requires local recalibration because MRKR has a substantially higher short-term arthroplasty event burden.",
    )

    add_heading(doc, "Introduction", 1)
    add_para(
        doc,
        "Knee osteoarthritis is heterogeneous: symptoms, radiographic severity, and treatment decisions often evolve on different time scales. A patient with stable pain but rapidly worsening joint-space narrowing may not have the same risk trajectory as a patient with severe pain and little structural progression. This mismatch creates a modeling problem. Static baseline models can be clinically convenient, but they may underuse the longitudinal information that prompts clinicians and patients to reconsider surgical timing.",
    )
    add_para(
        doc,
        "Public OA resources have made risk-model development more transparent. The OAI provides deeply phenotyped longitudinal symptoms, radiographs, and adjudicated knee replacement follow-up. Recent public-cohort studies have shown that imaging and clinical variables can predict OA progression and knee replacement, including multimodal radiograph models and MRI-based deep learning models. However, these models are often evaluated within similar research-cohort structures. Translation to real-world radiograph/EHR datasets requires a different question: not whether every OAI variable can be reproduced exactly, but whether a deliberately common-variable score retains useful rank ordering and can be recalibrated to local event rates.",
    )
    add_para(
        doc,
        "MRKR creates a rare opportunity for this transport question. It links a large, diverse real-world knee radiograph resource to pain scores, inferred KL grades, laterality metadata, diagnostic/procedural codes, and arthroplasty hardware. These features do not recreate OAI. They instead test whether an OAI-derived clinical-radiographic risk signal can survive measurement shifts, outcome ascertainment differences, and real-world case mix. We therefore separated the scientific OAI model from the transport model: Model E is the complete OAI landmark model, while Model F-core is an intentionally constrained OAI-MRKR common-variable model.",
    )

    add_heading(doc, "Methods", 1)
    add_para(
        doc,
        "Study design: This was a retrospective landmark prediction-model study with model development and internal validation in OAI, followed by transport validation and recalibration in MRKR. The OAI landmark was set at 24 months. Predictors measured at baseline and changes from baseline to the landmark were used to predict subsequent knee replacement. Knees with replacement before or at the landmark were excluded from post-landmark risk modeling.",
    )
    add_para(
        doc,
        "Cohort assembly: OAI knee rows were assembled from bilateral participant-level clinical, radiographic, and outcome files. We sequentially excluded knees with baseline/prior knee replacement, target-knee KR/TKA before or at the 24-month landmark, no positive post-landmark follow-up, missing landmark pain, and missing Model F-core predictors. MRKR knee-level baseline-landmark pairs were derived from side-specific radiographs without arthroplasty hardware at baseline; pairs were excluded for no positive hardware follow-up, missing or ambiguous laterality, missing landmark pain, missing baseline or landmark KL inference, or missing age/sex. The detailed counts are reported in Supplementary Tables 1 and 2.",
    )
    add_para(
        doc,
        "OAI primary model: Model E retained the full scientific predictor set available in OAI: age, sex, body mass index, knee side, baseline WOMAC pain, function and stiffness, baseline KL grade, baseline medial and lateral joint-space narrowing, and 0-24 month changes in WOMAC pain/function plus KL and medial/lateral joint-space narrowing. Cox proportional hazards models used knee-level observations with participant-level clustering where appropriate.",
    )
    add_para(
        doc,
        "OAI-MRKR transport model: Model F-core was fit in OAI using variables that could be approximately mapped in MRKR: age, sex, knee side, pain at landmark standardized within cohort, baseline KL grade, and KL worsening between baseline and landmark. MRKR laterality was harmonized to left/right knee-level observations. MRKR pain was not equivalent to WOMAC pain, so it was standardized rather than treated as directly interchangeable. MRKR KL grades were model-inferred, and KL worsening was defined as landmark inferred KL grade minus baseline inferred KL grade.",
    )
    add_para(
        doc,
        "Outcome definitions: In OAI, the primary outcome was target-knee knee replacement/total knee arthroplasty after the 24-month landmark, using the OAI outcomes file event date; follow-up time was event month minus 24 months for events and last outcome contact month minus 24 months for censored knees. In MRKR, the primary outcome was side-specific arthroplasty hardware first observed on a follow-up target-knee radiograph after the landmark. CPT-based knee arthroplasty outcomes were evaluated as sensitivity analyses, including patient-level CPT and combined hardware-or-CPT definitions. CPT was not used as the primary endpoint because many arthroplasty CPT rows were patient-level or lacked reliable left/right modifiers, whereas the prediction task and MRKR hardware endpoint were knee-side specific.",
    )
    add_para(
        doc,
        "MRKR algorithmic labels: MRKR image metadata were generated by the dataset authors using a ConvNeXt classifier trained on 6,000 labeled radiographs, with reported weighted F1-scores of 0.985 for laterality, 0.974 for view position, 0.981 for weight bearing, and 0.992 for arthroplasty. KL grades were inferred using an open-source Duke model with reported 75.9% exact accuracy in MOST and 93.5% of errors within one KL grade. The MRKR report explicitly states that no systematic manual validation of KL accuracy was performed within MRKR, so KL inference and hardware-defined outcome ascertainment were treated as real-world algorithmic labels and addressed in sensitivity analyses and limitations.",
    )
    add_para(
        doc,
        "Performance assessment: We evaluated Harrell C-index, time-dependent AUC, Brier score, mean predicted risk, Kaplan-Meier observed risk, calibration plots, decision-curve analysis, and risk strata. OAI internal validation used bootstrap optimism correction and repeated patient-level splits. Bootstrap 95% confidence intervals for Model F-core C-index and time-dependent AUC used 500 patient-level resamples within each cohort. MRKR transport validation reported original OAI-derived predictions and recalibrated predictions. Recalibration included slope plus baseline recalibration as the preferred method and offset-only recalibration as a diagnostic comparison. Strict MRKR sensitivity analyses excluded early events within 3, 6, and 12 months after the landmark.",
    )
    add_para(
        doc,
        "Proportional hazards assessment: For OAI Model E, Schoenfeld residual tests were used to evaluate proportional hazards globally and for each covariate. Because several terms showed evidence of time-varying effects, we fit a prespecified sensitivity model adding log(time) interactions for age, baseline pain, function change, baseline KL, KL change, and baseline medial JSN, and compared discrimination with the original Cox model.",
    )

    add_heading(doc, "Results", 1)
    add_heading(doc, "Cohort assembly and baseline differences", 2)
    add_para(
        doc,
        "The OAI and MRKR transport cohorts differed in ways directly relevant to transportability. MRKR had more women, substantially higher landmark pain, more KL worsening, shorter follow-up, more early arthroplasty hardware events, and a more racially diverse case mix. These differences support interpreting MRKR as a real-world transport validation and recalibration setting rather than a direct replication cohort.",
    )
    rows = [["Characteristic", "OAI", "MRKR", "SMD"]]
    for _, row in t["table1"].iterrows():
        oai_val = "" if pd.isna(row["OAI Model F-core (n=3104 knees)"]) else row["OAI Model F-core (n=3104 knees)"]
        mrkr_val = "" if pd.isna(row["MRKR validation (n=3412 knees)"]) else row["MRKR validation (n=3412 knees)"]
        rows.append(
            [
                row["Characteristic"],
                oai_val,
                mrkr_val,
                "" if pd.isna(row["SMD"]) else row["SMD"],
            ]
        )
    add_caption(doc, "Table 1. Baseline characteristics of OAI Model F-core training knees and MRKR validation knees.")
    add_table(doc, rows, widths=[2.5, 1.45, 1.45, 0.55])
    add_note(doc, "SMD denotes standardized mean difference. Race and ethnicity reflect source-data categories harmonized for reporting; MRKR race/ethnicity diversity is central to the transportability question.")

    add_heading(doc, "OAI dynamic landmark model", 2)
    add_para(
        doc,
        f"The OAI common-complete Model E analysis included 3,066 knees from 1,640 participants with 559 post-landmark knee replacement events. Model E had the strongest overall performance among the staged models, with apparent C-index {num(e['c_index'])}, optimism-corrected C-index {num(e['optimism_corrected_c_index'])}, and repeated split C-index {num(e['split_c_index_mean'])}. At 60 months, Model E had AUC {num(e['auc_60m'])}, Brier score {num(e['brier_60m'])}, mean predicted risk {pct(e['mean_pred_risk_60m'])}, and observed Kaplan-Meier risk {pct(e['observed_km_risk_60m'])}.",
    )
    add_para(
        doc,
        f"Compared with Model C, which included baseline symptoms and imaging, Model E provided substantial incremental value: likelihood-ratio chi-square {num(inc['lr_chisq'], 1)} on {int(inc['df'])} degrees of freedom (p={pval(inc['p_value'])}), delta C-index {num(inc['delta_c_index'])}, delta fixed-horizon 60-month AUC {num(inc['delta_auc_60m_fixed_logistic'])}, and delta AIC {num(inc['delta_aic'], 1)}.",
    )

    rows = [["Model", "Predictor set", "N knees", "Events", "C-index", "60-mo AUC", "60-mo Brier", "Optimism-corrected C"]]
    predictor_names = {
        "A_basic": "Age, sex, BMI, side",
        "B_baseline_symptoms": "A + baseline symptoms",
        "C_baseline_symptoms_imaging": "B + baseline KL/JSN",
        "D_dynamic_clinical": "B + symptom change",
        "E_dynamic_clinical_imaging": "C + symptom and structure change",
    }
    for _, row in t["model_cmp"].iterrows():
        rows.append(
            [
                row["model"].split("_")[0],
                predictor_names[row["model"]],
                str(int(row["n"])),
                str(int(row["events"])),
                num(row["c_index"]),
                num(row["auc_60m"]),
                num(row["brier_60m"]),
                num(row["optimism_corrected_c_index"]),
            ]
        )
    add_caption(doc, "Table 2. Staged OAI model performance in the 24-month landmark common-complete analysis set.")
    add_table(doc, rows, widths=[0.55, 2.2, 0.75, 0.65, 0.65, 0.7, 0.75, 0.95])
    add_note(doc, "Abbreviations: BMI, body mass index; JSN, joint-space narrowing; KL, Kellgren-Lawrence grade.")

    e_terms = [
        "pain_change",
        "kl_0",
        "jsn_medial_change",
        "jsn_lateral_change",
        "bmi",
    ]
    e_coef = t["coeff_e"].query("model == 'E_dynamic_clinical_imaging' and term in @e_terms")
    label = {
        "pain_change": "0-24 month WOMAC pain change",
        "kl_0": "Baseline KL grade",
        "jsn_medial_change": "0-24 month medial JSN change",
        "jsn_lateral_change": "0-24 month lateral JSN change",
        "bmi": "BMI",
    }
    rows = [["Predictor", "Hazard ratio", "95% CI", "P value"]]
    for _, row in e_coef.iterrows():
        rows.append(
            [
                label[row["term"]],
                num(row["hazard_ratio"]),
                f"{num(row['ci_lower_95'])}-{num(row['ci_upper_95'])}",
                pval(row["p_value"]),
            ]
        )
    add_caption(doc, "Table 3. Selected OAI Model E effects emphasizing dynamic symptom and structural information.")
    add_table(doc, rows, widths=[3.1, 1.0, 1.25, 0.85])

    add_heading(doc, "OAI-MRKR transport validation", 2)
    oai_time = t["time"].query("cohort == 'OAI'").iloc[0]
    mrkr_time = t["time"].query("cohort == 'MRKR'").iloc[0]
    add_para(
        doc,
        f"Model F-core used 3,104 OAI knees from 1,656 participants and 3,412 MRKR knees from 2,179 patients. MRKR had a markedly different time and event structure: {pct(mrkr_time['event_rate'])} of knees had primary hardware-defined events, median follow-up was {num(mrkr_time['median_followup_months'], 1)} months, and median event time was {num(mrkr_time['median_event_time_months'], 1)} months. In OAI, the event rate was {pct(oai_time['event_rate'])}, median follow-up was {num(oai_time['median_followup_months'], 1)} months, and median event time was {num(oai_time['median_event_time_months'], 1)} months.",
    )
    rows = [["Cohort", "Knees", "Patients", "Events", "Event rate", "Median follow-up, mo", "Events by 24 mo", "At risk at 24 mo"]]
    for _, row in t["time"].iterrows():
        rows.append(
            [
                row["cohort"],
                str(int(row["knees"])),
                str(int(row["patients"])),
                str(int(row["total_events"])),
                pct(row["event_rate"]),
                num(row["median_followup_months"], 1),
                str(int(row["events_24m"])),
                str(int(row["at_risk_24m"])),
            ]
        )
    add_caption(doc, "Table 4. Time and event structure for OAI Model F-core training and MRKR validation.")
    add_table(doc, rows, widths=[0.7, 0.65, 0.65, 0.65, 0.75, 1.2, 0.9, 0.9])

    doc.add_page_break()
    rows = [["Model F-core predictor", "Hazard ratio", "95% CI", "P value"]]
    labels_f = {
        "age": "Age",
        "female": "Female sex",
        "right_knee": "Right knee",
        "pain_landmark_z": "Standardized landmark pain",
        "kl_baseline": "Baseline KL grade",
        "kl_worsening": "KL worsening from baseline to landmark",
    }
    for _, row in t["fcoef"].iterrows():
        rows.append(
            [
                labels_f[row["term"]],
                num(row["hazard_ratio"]),
                f"{num(row['ci_lower_95'])}-{num(row['ci_upper_95'])}",
                pval(row["p_value"]),
            ]
        )
    add_caption(doc, "Table 5. OAI Model F-core coefficients for common-variable transport validation.")
    add_table(doc, rows, widths=[3.0, 1.0, 1.3, 0.9])

    rows = [["Cohort/model", "Horizon", "C-index", "AUC", "Brier", "Mean predicted", "Observed"]]
    keep = t["transport"][
        (t["transport"]["horizon_months"].isin([12, 24, 60]))
        & (
            (t["transport"]["cohort"] == "OAI")
            | (t["transport"]["model"].isin(["Original OAI-derived Model F-core", "MRKR slope+baseline recalibrated Model F-core"]))
        )
    ]
    for _, row in keep.iterrows():
        rows.append(
            [
                f"{row['cohort']}: {row['model'].replace(' Model F-core', '')}",
                f"{int(row['horizon_months'])} mo",
                num(row["c_index"]),
                num(row["auc_horizon"]),
                num(row["brier_horizon"]),
                pct(row["mean_predicted_risk"]),
                pct(row["observed_km_risk"]),
            ]
        )
    add_caption(doc, "Table 6. OAI apparent and MRKR transport validation performance across selected horizons.")
    add_table(doc, rows, widths=[2.35, 0.65, 0.65, 0.6, 0.6, 0.9, 0.75])
    add_note(doc, "MRKR recalibration used slope plus baseline recalibration and therefore changes absolute risk and Brier score, not ranking metrics.")

    add_para(
        doc,
        f"At 24 months, the original OAI-derived Model F-core substantially underestimated MRKR absolute risk ({pct(mrkr24['mean_predicted_risk'])} predicted vs {pct(mrkr24['observed_km_risk'])} observed). Slope plus baseline recalibration brought predicted risk close to observed risk ({pct(mrkr24_recal['mean_predicted_risk'])} vs {pct(mrkr24_recal['observed_km_risk'])}) and improved Brier score from {num(mrkr24['brier_horizon'])} to {num(mrkr24_recal['brier_horizon'])}.",
    )

    doc.add_page_break()
    add_heading(doc, "Risk strata and sensitivity analyses", 2)
    rows = [["Risk stratum", "Knees", "24-mo events", "Mean recalibrated risk", "Observed 24-mo risk"]]
    for _, row in t["strata"].iterrows():
        rows.append(
            [
                row["risk_group_recalibrated_24m"],
                str(int(row["n_knees"])),
                str(int(row["events_by_24m"])),
                pct(row["mean_recalibrated_predicted_24m_risk"]),
                pct(row["observed_km_24m_risk"]),
            ]
        )
    add_caption(doc, "Table 7. MRKR recalibrated 24-month risk strata.")
    add_table(doc, rows, widths=[1.25, 0.8, 0.9, 1.4, 1.25])

    rows = [["Sensitivity set", "Knees", "Events by 24 mo", "C-index", "AUC", "Observed risk", "Calibration slope"]]
    for _, row in t["strict"].iterrows():
        rows.append(
            [
                row["strict_definition"].replace("Exclude events <=", "Exclude <="),
                str(int(row["n_knees"])),
                str(int(row["events_by_24m"])),
                num(row["c_index"]),
                num(row["auc_24m"]),
                pct(row["observed_km_24m_risk"]),
                num(row["calibration_slope"]),
            ]
        )
    add_caption(doc, "Table 8. MRKR strict sensitivity analyses excluding early post-landmark events.")
    add_table(doc, rows, widths=[1.9, 0.7, 0.9, 0.6, 0.6, 0.9, 0.9])

    rows = [["Outcome definition", "Knees", "Total events", "24-mo events", "C-index", "AUC", "Observed 24-mo risk"]]
    outcome_labels = {
        "hardware_primary": "Side-specific hardware",
        "cpt_patient_level_sensitivity": "Patient-level CPT",
        "combined_hardware_or_cpt": "Hardware or CPT",
    }
    for _, row in t["outcome"].iterrows():
        rows.append(
            [
                outcome_labels.get(row["outcome_version"], row["outcome_version"]),
                str(int(row["n_knees"])),
                str(int(row["total_events"])),
                str(int(row["events_by_24m"])),
                num(row["c_index"]),
                num(row["auc_24m"]),
                pct(row["observed_km_24m_risk"]),
            ]
        )
    add_caption(doc, "Table 9. MRKR primary and sensitivity outcome definitions.")
    add_table(doc, rows, widths=[1.75, 0.65, 0.8, 0.8, 0.6, 0.6, 1.0])

    add_heading(doc, "Supplementary Analyses Added for Reviewer Robustness", 1)
    rows = [["Cohort/model", "Metric", "Horizon", "Estimate", "95% CI"]]
    for _, row in t["perf_ci"].iterrows():
        horizon = "Overall" if int(row["horizon_months"]) == 0 else f"{int(row['horizon_months'])} mo"
        metric = "C-index" if row["metric"] == "c_index" else "AUC"
        rows.append(
            [
                row["cohort_model"].replace(" Model F-core", ""),
                metric,
                horizon,
                num(row["estimate"]),
                f"{num(row['ci_lower_95'])}-{num(row['ci_upper_95'])}",
            ]
        )
    add_caption(doc, "Supplementary Table 1. Bootstrap confidence intervals for Model F-core discrimination.")
    add_table(doc, rows, widths=[2.4, 0.8, 0.8, 0.8, 1.1])
    add_note(doc, "Confidence intervals used 500 patient-level bootstrap resamples within each cohort.")

    add_caption(doc, "Supplementary Table 2. Reproducible Model F-core formula and baseline cumulative hazards.")
    betas = {row["term"]: row["beta"] for _, row in t["f_formula"].iterrows()}
    std_oai = t["f_std"].query("cohort == 'OAI training'").iloc[0]
    std_mrkr = t["f_std"].query("cohort == 'MRKR validation'").iloc[0]
    add_para(
        doc,
        "The OAI-derived Model F-core linear predictor is: LP = "
        f"{betas['age']:.6f}*age + {betas['female']:.6f}*female + {betas['right_knee']:.6f}*right_knee + "
        f"{betas['pain_landmark_z']:.6f}*pain_landmark_z + {betas['kl_baseline']:.6f}*baseline_KL + "
        f"{betas['kl_worsening']:.6f}*I(KL_change>=1). "
        f"In OAI, pain_landmark_z = (pain_0_10 - {std_oai['pain_landmark_mean']:.3f})/{std_oai['pain_landmark_sd']:.3f}; "
        f"in MRKR, pain_landmark_z = (pain_0_10 - {std_mrkr['pain_landmark_mean']:.3f})/{std_mrkr['pain_landmark_sd']:.3f}, reflecting cohort-standardized pain measurement rather than assuming WOMAC and MRKR pain scores are exchangeable on the raw scale.",
    )
    rows = [["Horizon", "OAI baseline cumulative hazard"]]
    for _, row in t["f_basehaz"].iterrows():
        rows.append([f"{int(row['horizon_months'])} mo", f"{row['baseline_cumulative_hazard']:.6f}"])
    add_table(doc, rows, widths=[1.2, 2.2])
    add_para(
        doc,
        "Original OAI-derived absolute risk at horizon t is computed as 1 - exp[-H0(t)*exp(LP)]. MRKR slope plus baseline recalibration fit Surv(time,event) ~ LP in MRKR. The recalibrated MRKR 24-month cumulative baseline hazard was "
        f"{t['recal_formula'].iloc[0]['baseline_cumulative_hazard_24m']:.6f}, and the calibration slope was {t['recal_formula'].iloc[0]['slope']:.3f}.",
    )

    rows = [["Step", "Knees", "Participants/patients", "Reason/detail"]]
    for _, row in t["oai_flow"].iterrows():
        rows.append([row["Step"], str(int(row["Knees"])), str(int(row["Participants/patients"])), "" if pd.isna(row["Reason/detail"]) else row["Reason/detail"]])
    add_caption(doc, "Supplementary Table 3. OAI inclusion and exclusion flow.")
    add_table(doc, rows, widths=[2.7, 0.7, 0.9, 2.0])

    rows = [["Step", "Knees", "Participants/patients", "Reason/detail"]]
    for _, row in t["mrkr_flow"].iterrows():
        rows.append([row["Step"], str(int(row["Knees"])), str(int(row["Participants/patients"])), "" if pd.isna(row["Reason/detail"]) else row["Reason/detail"]])
    add_caption(doc, "Supplementary Table 4. MRKR inclusion and exclusion flow.")
    add_table(doc, rows, widths=[2.7, 0.7, 0.9, 2.0])

    rows = [["Term", "Schoenfeld p value"]]
    for _, row in t["ph"].iterrows():
        if row["term"] in ["age", "pain_0", "function_change", "kl_0", "kl_change", "jsn_medial_0", "GLOBAL"]:
            rows.append([row["term"], pval(row["p_value"])])
    add_caption(doc, "Supplementary Table 5. Selected Schoenfeld residual proportional-hazards diagnostics for OAI Model E.")
    add_table(doc, rows, widths=[2.4, 1.2])

    rows = [["Model", "C-index", "60-mo AUC"]]
    for _, row in t["ph_sens"].iterrows():
        rows.append([row["model"], num(row["c_index"]), num(row["auc_60m"])])
    add_caption(doc, "Supplementary Table 6. OAI Model E proportional-hazards sensitivity analysis.")
    add_table(doc, rows, widths=[3.2, 0.9, 0.9])
    add_note(doc, "The time-varying coefficient sensitivity model added log(time) interactions for variables with prespecified or observed PH concerns. It preserved 60-month AUC but changed the global rank-order C-index modestly, supporting cautious interpretation of Cox coefficients and the emphasis on fixed-horizon validation metrics.")

    doc.add_page_break()
    add_heading(doc, "Figures", 1)
    fig_caps = [
        (
            "figure_oai_mrkr_sci_1_cohort_flow.png",
            "Figure 1. Cohort inclusion flow for OAI model development and MRKR transport validation. Flow diagram showing knee-level eligibility and exclusions for the Osteoarthritis Initiative (OAI) and the Multicenter Osteoarthritis Study Radiographic Knee Replacement (MRKR) dataset. OAI generated both the Model F-core training set and the Model E common-complete set; MRKR provided side-specific knee-level baseline-landmark pairs for real-world transport validation and recalibration.",
        ),
        (
            "figure_oai_mrkr_sci_2_modeling_framework.png",
            "Figure 2. Study design and modeling framework. Model E was developed in OAI as the full dynamic clinical-radiographic model. Model F-core was separately developed in OAI using variables approximately mappable between OAI and MRKR, then applied to MRKR for transport validation, recalibration, risk stratification, decision-curve analysis, and sensitivity analyses.",
        ),
        (
            "figure_oai_mrkr_sci_3_oai_staged_performance.png",
            "Figure 3. Staged OAI model performance in the 24-month landmark analysis. Sequential models show the incremental prognostic value of adding baseline symptoms, radiographic severity, and 0-24 month symptom and structural change, with Model E showing the strongest overall performance.",
        ),
        (
            "figure_oai_mrkr_sci_4_oai_model_e_calibration_dca.png",
            "Figure 4. Calibration and decision-curve analysis of OAI Model E at 60 months. The calibration plot compares mean predicted knee replacement risk with Kaplan-Meier observed risk across risk groups; decision-curve analysis estimates net benefit across threshold probabilities.",
        ),
        (
            "figure_oai_mrkr_sci_5_time_event_structure.png",
            "Figure 5. OAI and MRKR time-event structure. OAI had longer follow-up and later events, whereas MRKR had a higher and earlier arthroplasty burden. Numbers above bars indicate knees still at risk at each horizon.",
        ),
        (
            "figure_oai_mrkr_sci_6_mrkr_calibration_recalibration.png",
            "Figure 6. MRKR calibration before and after slope plus baseline recalibration. Original OAI-derived predictions retained moderate risk ranking but underestimated absolute MRKR risk; slope plus baseline recalibration shifted predicted probabilities toward observed Kaplan-Meier risks.",
        ),
        (
            "figure_oai_mrkr_sci_7_mrkr_risk_strata.png",
            "Figure 7. MRKR recalibrated 24-month risk strata. Knees were grouped into prespecified 24-month predicted-risk categories of <10%, 10-25%, 25-50%, and >50%; mean recalibrated risk and Kaplan-Meier observed risk were generally aligned across strata.",
        ),
        (
            "figure_oai_mrkr_sci_8_mrkr_sensitivity.png",
            "Figure 8. MRKR early-event exclusion and outcome-definition sensitivity analyses. The early-event analysis excluded events within 3, 6, and 12 months after landmark; the outcome analysis compared side-specific hardware, patient-level CPT, and combined hardware-or-CPT definitions.",
        ),
        (
            "supplementary_figure_s1_model_e_nomogram.png",
            "Supplementary Figure 1. Nomogram-style visualization for the full OAI dynamic Model E. This figure illustrates how baseline disease burden and early symptom-structural worsening contribute to predicted post-landmark knee replacement risk and is intended as a scientific visualization rather than a deployment-ready clinical calculator.",
        ),
        (
            "supplementary_figure_s2_model_f_calculator_mockup.png",
            "Supplementary Figure 2. Research-use web calculator interface for Model F-core. The prototype uses common OAI-MRKR variables and reports original OAI-derived and MRKR-recalibrated risk estimates; it is not intended for standalone clinical decision-making without prospective validation and local recalibration.",
        ),
        (
            "supplementary_figure_s3_mrkr_subgroups.png",
            "Supplementary Figure 3. MRKR subgroup performance at 24 months. AUC estimates are shown across clinically relevant demographic, pain, and radiographic subgroups and should be interpreted descriptively where sample size or event counts are limited.",
        ),
    ]
    for idx, (fname, cap) in enumerate(fig_caps):
        if idx > 0:
            doc.add_page_break()
        add_caption(doc, cap)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(FIGURES / fname), width=Inches(6.55))

    add_heading(doc, "Discussion", 1)
    add_para(
        doc,
        "This study supports three main conclusions. First, dynamic information matters. In OAI, the complete clinical-radiographic landmark model outperformed models that used only baseline symptoms, only baseline imaging, or only partial change information. Second, common-variable transport is feasible but imperfect. The OAI-derived Model F-core retained moderate discrimination in MRKR despite different pain measurement, inferred KL grades, real-world follow-up, and hardware-defined outcomes. Third, absolute risk requires local recalibration. The original OAI-derived score ranked MRKR knees reasonably but substantially underestimated MRKR short-term event risk because MRKR contained many earlier arthroplasty events and a higher baseline event burden.",
    )
    add_para(
        doc,
        "The MRKR results should therefore be interpreted as transport validation rather than direct external validation of the full OAI Model E. This distinction is a strength rather than a limitation. A direct Model E validation would require WOMAC domains and standardized OAI radiographic readings, which MRKR does not provide. By separating Model E from Model F-core, the analysis preserves the scientific value of the rich OAI model while asking a realistic deployment question: what survives when only clinically and radiographically mappable variables are available?",
    )
    add_para(
        doc,
        "The outcome-definition analyses further clarify the role of MRKR. Side-specific hardware detection is better aligned with knee-level prediction than patient-level CPT procedure codes, but hardware detection can lag actual surgery date and may capture prior procedures visible on follow-up imaging. CPT-based sensitivity analyses increased event counts but reduced discrimination, consistent with noise introduced by patient-level or non-side-specific coding. Early-event exclusions attenuated but did not eliminate transport performance, suggesting that the model signal was not driven solely by imminent pre-existing surgical scheduling.",
    )
    add_para(
        doc,
        "Race and ethnicity are not peripheral in this transport setting. MRKR contained a substantially larger proportion of Black patients than OAI and had higher landmark pain and earlier arthroplasty events. These differences may reflect referral patterns, disease severity, access to arthroplasty, imaging frequency, and health-system capture. We therefore report race/ethnicity distributions and subgroup performance rather than treating transportability as a purely technical measurement problem. The findings support recalibrated risk ranking in MRKR, but they do not establish that the same decision thresholds should be used across racial, ethnic, or health-system groups.",
    )
    add_para(
        doc,
        "These findings have practical implications for public-data OA modeling. Papers that report only AUC may overstate clinical readiness. Calibration, risk strata, decision curves, and local recalibration are essential when moving from a research cohort to a real-world imaging/EHR dataset. The recalibrated MRKR strata showed clinically interpretable separation, supporting potential use for cohort enrichment, trial planning, and shared decision support after local validation rather than immediate individual-level surgical recommendation.",
    )

    add_heading(doc, "Limitations", 1)
    add_para(
        doc,
        "Several limitations should be emphasized. MRKR pain scores are not WOMAC pain and were therefore standardized rather than directly equated. MRKR KL grades are model-inferred rather than centrally adjudicated radiographic grades. The MRKR report cited prior MOST performance for the KL inference model but did not perform systematic manual validation inside MRKR; this means KL measurement error could contribute to attenuated discrimination and calibration shift. MRKR hardware outcomes are side-specific and the source metadata classifier reported high arthroplasty F1, but hardware first observed on imaging may lag the true surgery date. CPT codes provide procedural information but often lack reliable laterality, which is why they were retained as sensitivity rather than primary outcomes. OAI Model E used a common-complete set for fair staged-model comparison; excluded knees had lower event rates and milder symptom burden, so generalizability beyond the complete-case subset requires additional missing-data work. Proportional hazards diagnostics showed global non-proportionality and several variable-specific departures; therefore, we added a time-varying coefficient sensitivity analysis and emphasized fixed-horizon AUC, calibration, and Brier score alongside Cox coefficients. Residual non-proportionality remains a reason to treat individual hazard ratios as time-averaged summaries rather than mechanistic constants.",
    )

    add_heading(doc, "Conclusions", 1)
    add_para(
        doc,
        "In OAI, a 24-month landmark model combining symptoms, radiographic severity, and longitudinal change improved prediction of subsequent knee replacement. A common-variable OAI transport model retained clinically meaningful ranking performance in MRKR but required recalibration to account for the higher real-world short-term arthroplasty burden. The analysis provides a transparent template for developing rich cohort models while testing practical transportability in public real-world imaging/EHR data.",
    )

    add_heading(doc, "Data and code availability", 1)
    add_para(
        doc,
        "OAI and MRKR are public or controlled-access research resources. The derived analysis datasets, data dictionaries, code scripts, tables, and figures used for this draft are stored in the local project folders derived/, scripts/, results/tables/, and results/figures/. Release-ready code should include a reproducible environment file and data-access instructions rather than redistributing restricted source data.",
    )

    add_heading(doc, "References", 1)
    refs = [
        "MRKR dataset preprint. Multicenter Osteoarthritis Study Radiographic Knee Replacement dataset. arXiv. https://arxiv.org/abs/2411.00866",
        "Tiulpin A, et al. Multimodal machine learning-based knee osteoarthritis progression prediction from plain radiographs and clinical data. arXiv. https://arxiv.org/abs/1904.06236",
        "Tolpadi AA, et al. Deep learning predicts total knee replacement from magnetic resonance images. Scientific Reports. https://www.nature.com/articles/s41598-020-63395-9",
        "Rajamohan HR, et al. Prediction of future knee replacement from baseline MRI data. Scientific Reports. https://www.nature.com/articles/s41598-023-33934-1",
        "Zhao Y, et al. Machine learning using OAI imaging and semiquantitative features for knee pain severity. Scientific Reports. https://www.nature.com/articles/s41598-024-65613-0",
        "Swiecicki A, et al. Deep learning-based algorithm for assessment of knee osteoarthritis severity in radiographs matches performance of radiologists. arXiv. https://arxiv.org/abs/2207.12521",
        "Kim B, et al. Automated detection of surgical implants on plain knee radiographs using a deep learning algorithm. Medicina. 2022;58:1677. https://doi.org/10.3390/medicina58111677",
        "Van Calster B, et al. Calibration: the Achilles heel of predictive analytics. BMC Medicine. https://link.springer.com/article/10.1186/s12916-019-1466-7",
        "Vickers AJ, et al. Interpreting decision curve analysis. Diagnostic and Prognostic Research. https://link.springer.com/article/10.1186/s41512-019-0064-7",
        "TRIPOD+AI reporting guidance for prediction models using artificial intelligence. BMJ. https://www.bmj.com/content/385/bmj-2023-078378",
    ]
    for i, ref in enumerate(refs, start=1):
        add_para(doc, f"{i}. {ref}")

    # Keep wide tables readable.
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT

    doc.save(DOCX_OUT)


def main() -> None:
    tables = load_tables()
    make_literature_synthesis(tables)
    build_docx(tables)
    print(DOCX_OUT)
    print(LIT_OUT)


if __name__ == "__main__":
    main()
