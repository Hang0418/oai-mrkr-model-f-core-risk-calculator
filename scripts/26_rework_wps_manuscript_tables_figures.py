from __future__ import annotations

from pathlib import Path
import math

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "results" / "tables"
FIGURES = ROOT / "results" / "figures"
SRC = ROOT / "drafts" / "OAI_MRKR_high_impact_manuscript_from_WPS_reworked_clean.docx"
OUT = ROOT / "drafts" / "OAI_MRKR_high_impact_manuscript_WPS_basis_tables_figures_reworked.docx"


def fnum(x, digits=3):
    if pd.isna(x):
        return ""
    return f"{float(x):.{digits}f}"


def fpct(x, digits=1):
    if pd.isna(x):
        return ""
    return f"{100 * float(x):.{digits}f}%"


def fp(x):
    if pd.isna(x):
        return ""
    x = float(x)
    if x < 0.001:
        return f"{x:.1e}"
    return f"{x:.3f}"


def clean_cell(x) -> str:
    if pd.isna(x):
        return ""
    s = str(x)
    return "" if s.lower() == "nan" else s


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.3)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        if name in styles:
            styles[name].font.name = "Arial"
            styles[name].font.color.rgb = RGBColor(31, 53, 85)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9E2EC")
        borders.append(tag)
    tbl_pr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, bold=False, size=7.8, align="left") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)


def note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(7.6)
    r.font.color.rgb = RGBColor(90, 90, 90)


def add_table(doc: Document, rows: list[list[str]], widths: list[float], numeric_cols: set[int] | None = None) -> None:
    numeric_cols = numeric_cols or set()
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, val in enumerate(rows[0]):
        set_cell_shading(hdr.cells[i], "EAF1F8")
        set_cell_text(hdr.cells[i], val, bold=True, size=7.8, align="center" if i in numeric_cols else "left")
        hdr.cells[i].width = Inches(widths[i])
    for values in rows[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(values):
            set_cell_text(cells[i], val, size=7.6, align="center" if i in numeric_cols else "left")
            cells[i].width = Inches(widths[i])
    set_table_borders(table)


def add_fig(doc: Document, stem: str, width: float) -> None:
    path = FIGURES / f"{stem}.png"
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def copy_front_and_refs(src: Document) -> tuple[list[str], list[str]]:
    texts = [p.text for p in src.paragraphs]
    table_start = next(i for i, t in enumerate(texts) if t.strip().startswith("Table 1. Baseline characteristics"))
    refs_start = next(i for i, t in enumerate(texts) if t.strip() == "References")
    front = texts[:table_start]
    refs = texts[refs_start:]
    return front, refs


def update_crossrefs(text: str) -> str:
    replacements = [
        ("The overall study design and analysis workflow are shown in Figure 1.", "The cohort flow and modeling framework are shown in Figure 1 and Figure 2."),
        ("Staged model performance is summarized in Table 1 and Figure 2.", "Baseline cohort differences are summarized in Table 1. Staged model performance is summarized in Table 2 and Figure 3."),
        ("Selected Model E effects are shown in Table 2.", "Selected Model E effects are shown in Table 3."),
        ("Figure 3.", "Figure 4."),
        ("Table 3 and Figure 4.", "Table 4 and Figure 5."),
        ("The OAI-derived Model F-core coefficients are shown in Table 4.", "The OAI-derived Model F-core coefficients are shown in Table 5."),
        ("Transport validation performance across prediction horizons is shown in Table 5, and calibration before and after recalibration is shown in Figure 5.", "Transport validation performance across prediction horizons is shown in Table 6, and calibration before and after recalibration is shown in Figure 6."),
        ("summarized in Table 6 and Figure 6.", "summarized in Table 7 and Figure 7."),
        ("summarized in Table 7, Table 8, and Figure 7.", "summarized in Table 8, Table 9, and Figure 8."),
        ("shown in Supplementary Figure 1.", "shown in Supplementary Figure 3."),
        (
            "To assess potential selection bias, included and excluded knees should be compared in supplementary analyses using baseline demographic, symptom, radiographic, and outcome characteristics.",
            "To assess potential selection bias, included and excluded knees were compared in Supplementary Tables 7 and 8 using demographic, symptom, radiographic, follow-up, and outcome characteristics.",
        ),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def add_plain_para(doc: Document, text: str) -> None:
    if not text.strip():
        doc.add_paragraph("")
        return
    if text.strip() in {"Abstract", "Introduction", "Methods", "Results", "Discussion", "References"}:
        p = doc.add_heading(text.strip(), level=1)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.color.rgb = RGBColor(31, 53, 85)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.06
    r = p.add_run(update_crossrefs(text))
    r.font.name = "Arial"
    r.font.size = Pt(10.3)


def table1(doc):
    df = pd.read_csv(TABLES / "oai_mrkr_reviewer_table1_baseline_smd.csv")
    rows = [["Characteristic", "OAI Model F-core", "MRKR validation", "SMD"]]
    char_col, oai_col, mrkr_col, smd_col = df.columns[:4]
    for _, r in df.iterrows():
        rows.append([clean_cell(r[char_col]), clean_cell(r[oai_col]), clean_cell(r[mrkr_col]), fnum(r[smd_col], 3) if not pd.isna(r[smd_col]) else ""])
    caption(doc, "Table 1. Baseline characteristics of OAI Model F-core training knees and MRKR validation knees.")
    add_table(doc, rows, [2.75, 1.65, 1.65, 0.75], {3})
    note(doc, "SMD, standardized mean difference; KL, Kellgren-Lawrence grade; IQR, interquartile range.")


def table2(doc):
    df = pd.read_csv(TABLES / "oai_docx_plan_model_comparison.csv")
    rows = [["Model", "N knees", "Events", "Predictors", "C-index", "Optimism-corrected C", "Split C", "60-mo AUC", "60-mo Brier", "AIC"]]
    labels = {
        "A_basic": "A",
        "B_baseline_symptoms": "B",
        "C_baseline_symptoms_imaging": "C",
        "D_dynamic_clinical": "D",
        "E_dynamic_clinical_imaging": "E",
    }
    pred = {
        "A_basic": "Age, sex, BMI, side",
        "B_baseline_symptoms": "+ baseline WOMAC",
        "C_baseline_symptoms_imaging": "+ baseline KL/JSN",
        "D_dynamic_clinical": "+ symptom change",
        "E_dynamic_clinical_imaging": "+ symptom and structural change",
    }
    for _, r in df.iterrows():
        rows.append([labels[r["model"]], str(int(r["n"])), str(int(r["events"])), pred[r["model"]], fnum(r["c_index"]), fnum(r["optimism_corrected_c_index"]), fnum(r["split_c_index_mean"]), fnum(r["auc_60m"]), fnum(r["brier_60m"]), fnum(r["aic"], 1)])
    caption(doc, "Table 2. Staged OAI model performance in the 24-month landmark common-complete analysis set.")
    add_table(doc, rows, [0.42, 0.55, 0.55, 1.35, 0.58, 0.85, 0.58, 0.65, 0.65, 0.58], {1, 2, 4, 5, 6, 7, 8, 9})


def table3(doc):
    df = pd.read_csv(TABLES / "oai_docx_plan_cox_coefficients.csv")
    keep = ["kl_0", "pain_change", "jsn_medial_change", "jsn_lateral_change", "bmi", "female"]
    labels = {
        "kl_0": "Baseline KL grade",
        "pain_change": "0–24 month WOMAC pain change",
        "jsn_medial_change": "0–24 month medial JSN change",
        "jsn_lateral_change": "0–24 month lateral JSN change",
        "bmi": "BMI",
        "female": "Female sex",
    }
    df = df[(df["model"] == "E_dynamic_clinical_imaging") & (df["term"].isin(keep))].copy()
    df["ord"] = df["term"].map({k: i for i, k in enumerate(keep)})
    df = df.sort_values("ord")
    rows = [["Predictor", "Hazard ratio", "95% CI", "P value"]]
    for _, r in df.iterrows():
        rows.append([labels[r["term"]], fnum(r["hazard_ratio"]), f"{fnum(r['ci_lower_95'])}-{fnum(r['ci_upper_95'])}", fp(r["p_value"])])
    caption(doc, "Table 3. Selected OAI Model E effects emphasizing dynamic symptom and structural information.")
    add_table(doc, rows, [2.7, 1.1, 1.3, 1.0], {1, 2, 3})


def table4(doc):
    df = pd.read_csv(TABLES / "oai_mrkr_plan_time_structure_table.csv")
    rows = [["Cohort", "Knees", "Patients", "Total events", "Event rate", "Median follow-up", "Median event time", "12-mo events", "24-mo events", "At risk at 24 mo"]]
    for _, r in df.iterrows():
        rows.append([r["cohort"], str(int(r["knees"])), str(int(r["patients"])), str(int(r["total_events"])), fpct(r["event_rate"]), f"{fnum(r['median_followup_months'],1)} mo", f"{fnum(r['median_event_time_months'],1)} mo", str(int(r["events_12m"])), str(int(r["events_24m"])), str(int(r["at_risk_24m"]))])
    caption(doc, "Table 4. Time and event structure for OAI Model F-core training and MRKR validation.")
    add_table(doc, rows, [0.55, 0.55, 0.55, 0.7, 0.65, 0.8, 0.8, 0.65, 0.65, 0.85], set(range(1, 10)))


def table5(doc):
    df = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_f_formula_coefficients.csv")
    labels = {
        "age": "Age, years",
        "female": "Female sex",
        "right_knee": "Right knee",
        "pain_landmark_z": "Standardized landmark pain",
        "kl_baseline": "Baseline KL grade",
        "kl_worsening": "KL worsening",
    }
    rows = [["Predictor", "Beta", "Hazard ratio", "95% CI", "P value"]]
    for _, r in df.iterrows():
        rows.append([labels[r["term"]], fnum(r["beta"], 6), fnum(r["hazard_ratio"]), f"{fnum(r['ci_lower_95'])}-{fnum(r['ci_upper_95'])}", fp(r["p_value"])])
    caption(doc, "Table 5. OAI Model F-core coefficients for common-variable transport validation.")
    add_table(doc, rows, [2.25, 1.0, 1.0, 1.35, 1.0], {1, 2, 3, 4})


def table6(doc):
    df = pd.read_csv(TABLES / "oai_mrkr_plan_transport_metrics_by_horizon.csv")
    rows = [["Cohort/model", "Horizon", "C-index", "AUC", "Mean predicted", "Observed", "Brier"]]
    for _, r in df.iterrows():
        label = r["cohort"] + " " + r["model"].replace("MRKR ", "").replace(" Model F-core", "")
        rows.append([label, f"{int(r['horizon_months'])} mo", fnum(r["c_index"]), fnum(r["auc_horizon"]), fpct(r["mean_predicted_risk"]), fpct(r["observed_km_risk"]), fnum(r["brier_horizon"])])
    caption(doc, "Table 6. OAI apparent and MRKR transport validation performance of Model F-core before and after recalibration.")
    add_table(doc, rows, [2.45, 0.65, 0.75, 0.7, 0.95, 0.75, 0.65], set(range(1, 7)))
    note(doc, "The original OAI-derived model was applied to MRKR without refitting. Slope plus baseline recalibration updated absolute risk estimates while preserving rank ordering; therefore, C-index and AUC were unchanged after recalibration. Offset-only baseline recalibration updates the baseline risk while preserving the OAI-derived calibration slope; it is shown as a diagnostic comparison rather than the preferred recalibration approach.")


def table7(doc):
    df = pd.read_csv(TABLES / "oai_mrkr_plan_mrkr_risk_strata_24m.csv")
    labels = {"<10%": "Low", "10-25%": "Intermediate", "25-50%": "High", ">50%": "Very high"}
    rows = [["Risk label", "Risk stratum", "Knees", "24-mo events", "Mean recalibrated risk", "Observed 24-mo risk"]]
    for _, r in df.iterrows():
        risk_group = r["risk_group_recalibrated_24m"]
        rows.append([labels.get(risk_group, ""), risk_group, str(int(r["n_knees"])), str(int(r["events_by_24m"])), fpct(r["mean_recalibrated_predicted_24m_risk"]), fpct(r["observed_km_24m_risk"])])
    caption(doc, "Table 7. MRKR recalibrated 24-month risk strata.")
    add_table(doc, rows, [0.95, 1.0, 0.7, 0.85, 1.45, 1.25], {2, 3, 4, 5})


def table8(doc):
    df = pd.read_csv(TABLES / "oai_mrkr_plan_mrkr_strict_sensitivity_24m.csv")
    rows = [["Sensitivity set", "Knees", "Events by 24 mo", "C-index", "AUC", "Observed risk", "Calibration slope"]]
    for _, r in df.iterrows():
        rows.append([r["strict_definition"], str(int(r["n_knees"])), str(int(r["events_by_24m"])), fnum(r["c_index"]), fnum(r["auc_24m"]), fpct(r["observed_km_24m_risk"]), fnum(r["calibration_slope"])])
    caption(doc, "Table 8. MRKR strict sensitivity analyses excluding early post-landmark events.")
    add_table(doc, rows, [1.45, 0.7, 1.0, 0.75, 0.7, 0.9, 1.0], set(range(1, 7)))


def table9(doc):
    df = pd.read_csv(TABLES / "oai_mrkr_plan_mrkr_outcome_sensitivity_24m.csv")
    rows = [["Outcome definition", "Knees", "Total events", "24-mo events", "C-index", "AUC", "Observed 24-mo risk"]]
    for _, r in df.iterrows():
        label = {
            "hardware_primary": "Side-specific hardware",
            "cpt_patient_level_sensitivity": "Patient-level CPT",
            "combined_hardware_or_cpt": "Hardware or CPT",
        }.get(r["outcome_version"], r["outcome_version"])
        rows.append([label, str(int(r["n_knees"])), str(int(r["total_events"])), str(int(r["events_by_24m"])), fnum(r["c_index"]), fnum(r["auc_24m"]), fpct(r["observed_km_24m_risk"])])
    caption(doc, "Table 9. MRKR primary and sensitivity outcome definitions.")
    add_table(doc, rows, [1.45, 0.7, 0.85, 0.9, 0.7, 0.7, 1.0], set(range(1, 7)))


def supp_tables(doc):
    doc.add_page_break()
    h = doc.add_heading("Supplementary Analyses", level=1)
    for r in h.runs:
        r.font.name = "Arial"
        r.font.color.rgb = RGBColor(31, 53, 85)

    for title, filename in [
        ("Supplementary Table 1. OAI inclusion and exclusion flow.", "oai_reviewer_inclusion_exclusion_flow.csv"),
        ("Supplementary Table 2. MRKR inclusion and exclusion flow.", "mrkr_reviewer_inclusion_exclusion_flow.csv"),
    ]:
        df = pd.read_csv(TABLES / filename)
        rows = [["Step", "Knees", "Participants/patients", "Reason/detail"]]
        for _, r in df.iterrows():
            rows.append([r["Step"], str(int(r["Knees"])), str(int(r["Participants/patients"])), "" if pd.isna(r["Reason/detail"]) else r["Reason/detail"]])
        caption(doc, title)
        add_table(doc, rows, [2.35, 0.75, 1.25, 2.2], {1, 2})

    ci = pd.read_csv(TABLES / "oai_mrkr_reviewer_discrimination_ci.csv")
    rows = [["Cohort/model", "Metric", "Horizon", "Estimate", "95% CI"]]
    for _, r in ci.iterrows():
        horizon = "Overall" if int(r["horizon_months"]) == 0 else f"{int(r['horizon_months'])} mo"
        rows.append([r["cohort_model"], r["metric"].upper() if r["metric"] == "auc" else "C-index", horizon, fnum(r["estimate"]), f"{fnum(r['ci_lower_95'])}-{fnum(r['ci_upper_95'])}"])
    caption(doc, "Supplementary Table 3. Bootstrap confidence intervals for Model F-core discrimination.")
    add_table(doc, rows, [2.25, 0.9, 0.8, 0.85, 1.1], {1, 2, 3, 4})
    note(doc, "Confidence intervals used 500 patient-level bootstrap resamples within each cohort.")

    coef = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_f_formula_coefficients.csv")
    beta = dict(zip(coef["term"], coef["beta"]))
    std = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_f_standardization.csv")
    recal = pd.read_csv(TABLES / "oai_mrkr_reviewer_recalibration_formula.csv").iloc[0]
    base = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_f_baseline_hazard.csv")
    caption(doc, "Supplementary Table 4. Reproducible Model F-core formula and baseline cumulative hazards.")
    formula = (
        f"Linear predictor: LP = {beta['age']:.6f}*age + {beta['female']:.6f}*female + "
        f"{beta['right_knee']:.6f}*right_knee + {beta['pain_landmark_z']:.6f}*pain_landmark_z + "
        f"{beta['kl_baseline']:.6f}*baseline_KL + {beta['kl_worsening']:.6f}*I(KL_change>=1). "
        f"OAI pain z = (pain_0_10 - {std.iloc[0]['pain_landmark_mean']:.3f})/{std.iloc[0]['pain_landmark_sd']:.3f}; "
        f"MRKR pain z = (pain_0_10 - {std.iloc[1]['pain_landmark_mean']:.3f})/{std.iloc[1]['pain_landmark_sd']:.3f}. "
        f"Original risk = 1 - exp[-H0(t)*exp(LP)]. MRKR recalibration slope = {recal['slope']:.3f}; "
        f"MRKR recalibrated 24-month baseline cumulative hazard = {recal['baseline_cumulative_hazard_24m']:.6f}."
    )
    note(doc, formula)
    rows = [["Horizon", "OAI baseline cumulative hazard"]]
    for _, r in base.iterrows():
        rows.append([f"{int(r['horizon_months'])} mo", f"{r['baseline_cumulative_hazard']:.6f}"])
    add_table(doc, rows, [1.4, 2.1], {1})

    ph = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_e_schoenfeld_ph.csv")
    rows = [["Term", "Schoenfeld p value"]]
    for _, r in ph.iterrows():
        rows.append([r["term"], fp(r["p_value"])])
    caption(doc, "Supplementary Table 5. Selected Schoenfeld residual proportional-hazards diagnostics for OAI Model E.")
    add_table(doc, rows, [2.4, 1.6], {1})

    phs = pd.read_csv(TABLES / "oai_mrkr_reviewer_model_e_ph_sensitivity_performance.csv")
    rows = [["Model", "C-index", "60-mo AUC"]]
    for _, r in phs.iterrows():
        rows.append([r["model"], fnum(r["c_index"]), fnum(r["auc_60m"])])
    caption(doc, "Supplementary Table 6. OAI Model E proportional-hazards sensitivity analysis.")
    add_table(doc, rows, [3.0, 1.0, 1.0], {1, 2})
    note(doc, "The time-varying coefficient sensitivity model added log(time) interactions for variables with prespecified or observed PH concerns.")

    for title, filename in [
        ("Supplementary Table 7. OAI Model F-core included versus excluded comparison.", "oai_model_f_included_vs_excluded_comparison.csv"),
        ("Supplementary Table 8. MRKR Model F-core included versus excluded comparison.", "mrkr_model_f_included_vs_excluded_comparison.csv"),
    ]:
        df = pd.read_csv(TABLES / filename)
        rows = [["Characteristic", "Included", "Excluded", "Included non-missing", "Excluded non-missing", "SMD"]]
        for _, r in df.iterrows():
            rows.append([
                clean_cell(r["Characteristic"]),
                clean_cell(r["Included"]),
                clean_cell(r["Excluded"]),
                clean_cell(r["Included non-missing"]),
                clean_cell(r["Excluded non-missing"]),
                clean_cell(r["SMD"]),
            ])
        caption(doc, title)
        add_table(doc, rows, [2.35, 1.05, 1.05, 0.8, 0.8, 0.55], {1, 2, 3, 4, 5})
        note(doc, "Included knees were complete for Model F-core variables; excluded knees were incomplete or otherwise unavailable for the mapped Model F-core analysis. SMD values >=0.10 indicate potentially meaningful imbalance.")


def add_figures(doc):
    doc.add_page_break()
    h = doc.add_heading("Figures", level=1)
    for r in h.runs:
        r.font.name = "Arial"
        r.font.color.rgb = RGBColor(31, 53, 85)
    figs = [
        ("Figure 1. Cohort inclusion flow for OAI model development and MRKR transport validation.", "figure_oai_mrkr_sci_1_cohort_flow", 6.5),
        ("Figure 2. Study design and modeling framework.", "figure_oai_mrkr_sci_2_modeling_framework", 6.5),
        ("Figure 3. Staged OAI model performance in the 24-month landmark analysis.", "figure_oai_mrkr_sci_3_oai_staged_performance", 6.5),
        ("Figure 4. Calibration and decision-curve analysis of OAI Model E at 60 months.", "figure_oai_mrkr_sci_4_oai_model_e_calibration_dca", 6.5),
        ("Figure 5. OAI and MRKR time-event structure.", "figure_oai_mrkr_sci_5_time_event_structure", 6.2),
        ("Figure 6. MRKR transport validation and recalibration of the OAI-derived Model F-core.", "figure_oai_mrkr_sci_6_mrkr_calibration_recalibration", 6.3),
        ("Figure 7. MRKR recalibrated 24-month risk strata.", "figure_oai_mrkr_sci_7_mrkr_risk_strata", 6.0),
        ("Figure 8. MRKR early-event exclusion and outcome-definition sensitivity analyses.", "figure_oai_mrkr_sci_8_mrkr_sensitivity", 6.3),
        ("Supplementary Figure 1. Formal nomogram for the full OAI dynamic Model E.", "supplementary_figure_s1_model_e_nomogram", 6.3),
        ("Supplementary Figure 2. Model F-core research-use risk calculator.", "supplementary_figure_s2_model_f_calculator_mockup", 6.3),
        ("Supplementary Figure 3. MRKR subgroup performance at 24 months.", "supplementary_figure_s3_mrkr_subgroups", 5.9),
    ]
    for title, stem, width in figs:
        caption(doc, title)
        add_fig(doc, stem, width)


def build() -> None:
    src = Document(SRC)
    front, refs = copy_front_and_refs(src)
    doc = Document()
    style_doc(doc)
    for text in front:
        add_plain_para(doc, text)

    h = doc.add_heading("Tables", level=1)
    for r in h.runs:
        r.font.name = "Arial"
        r.font.color.rgb = RGBColor(31, 53, 85)
    for fn in (table1, table2, table3, table4, table5, table6, table7, table8, table9):
        fn(doc)
    supp_tables(doc)
    add_figures(doc)
    doc.add_page_break()
    for text in refs:
        add_plain_para(doc, text)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
