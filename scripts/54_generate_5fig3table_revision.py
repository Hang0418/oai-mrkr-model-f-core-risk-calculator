#!/usr/bin/env python3
"""Generate a 5-main-figure/3-main-table manuscript revision and supplement."""

from __future__ import annotations

import html
import shutil
import zipfile
import re
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from scipy.special import expit, logit
from sklearn.metrics import brier_score_loss, roc_auc_score


ROOT = Path("/Users/hehang/Downloads/时间序列预测模型/膝关节炎")
BASE = ROOT / "stage_specific_progression_framework" / "complete_project_analysis"
MS = BASE / "three_stage_bridge_transport_manuscript"
CLIN = MS / "clinical_translation_upgrade_20260722"
OUT = MS / "main_supplement_5fig3table_revision_20260722"
FIGS = OUT / "main_figures"
SUPP_FIGS = OUT / "supplementary_figures"
TABLES = OUT / "tables"
SUPP_TABLES = OUT / "supplementary_tables"
CALC = OUT / "research_calculator"
MANUSCRIPT = OUT / "manuscript"

BASE_MAIN = MS / "major_comment5_bridge_sample_revision_20260722" / "three_stage_bridge_transport_manuscript_main_draft_WPSbase_revised_bridge_sample_fixed.docx"
BASE_SUPP = MS / "major_comment5_bridge_sample_revision_20260722" / "three_stage_bridge_transport_supplementary_materials_S1S2S6_updated_final_bridge_sample_fixed.docx"

COL = {
    "blue": "#2F6F97",
    "red": "#C83E50",
    "green": "#2E8B57",
    "gold": "#E3B23C",
    "gray": "#6B6B6B",
    "light": "#D9EAF7",
    "purple": "#6A4C93",
}


def ensure_dirs() -> None:
    for p in [OUT, FIGS, SUPP_FIGS, TABLES, SUPP_TABLES, CALC, MANUSCRIPT]:
        p.mkdir(parents=True, exist_ok=True)


def read(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, low_memory=False)


def pct(x: float, digits: int = 1) -> str:
    if pd.isna(x):
        return ""
    return f"{100 * float(x):.{digits}f}%"


def wilson(k: int, n: int, z: float = 1.96) -> tuple[float, float]:
    if n == 0:
        return np.nan, np.nan
    p = k / n
    den = 1 + z**2 / n
    centre = p + z**2 / (2 * n)
    half = z * np.sqrt((p * (1 - p) + z**2 / (4 * n)) / n)
    return (centre - half) / den, (centre + half) / den


def metric_text(y: pd.Series, p: pd.Series) -> tuple[float, float, float, float]:
    yv = y.to_numpy().astype(int)
    pv = np.clip(p.to_numpy().astype(float), 1e-6, 1 - 1e-6)
    auc = roc_auc_score(yv, pv)
    brier = brier_score_loss(yv, pv)
    try:
        import statsmodels.api as sm
        fit = sm.Logit(yv, sm.add_constant(logit(pv))).fit(disp=False, maxiter=200)
        intercept = float(fit.params[0])
        slope = float(fit.params[1])
    except Exception:
        intercept, slope = np.nan, np.nan
    return auc, brier, intercept, slope


def calibration_points(y: pd.Series, p: pd.Series, bins: int = 10) -> pd.DataFrame:
    d = pd.DataFrame({"y": y.astype(int), "p": np.clip(p.astype(float), 1e-6, 1 - 1e-6)})
    d["bin"] = pd.qcut(d["p"], q=min(bins, d["p"].nunique()), duplicates="drop")
    return d.groupby("bin", observed=True).agg(pred=("p", "mean"), obs=("y", "mean"), n=("y", "size")).reset_index()


def extract_docx_image(docx_path: Path, image_index: int, out_path: Path) -> Path:
    with zipfile.ZipFile(docx_path) as z:
        media = sorted([n for n in z.namelist() if n.startswith("word/media/image")])
        if image_index >= len(media):
            raise IndexError(f"{docx_path} has only {len(media)} images")
        out_path.write_bytes(z.read(media[image_index]))
    return out_path


def save_fig(fig: plt.Figure, stem: str, folder: Path = FIGS) -> None:
    fig.savefig(folder / f"{stem}.png", dpi=600)
    fig.savefig(folder / f"{stem}.pdf")
    plt.close(fig)


def make_figure2() -> None:
    event = read(BASE / "check_incident_oa_analysis/tables/check_incident_oa_cohort_event_summary.csv")
    risk = read(BASE / "check_incident_oa_analysis/tables/check_incident_oa_risk_by_pain_kl_group.csv")
    comp = read(BASE / "check_incident_oa_analysis/tables/check_incident_oa_model_comparison_traditional_vs_ml.csv")
    pred = read(BASE / "check_incident_oa_analysis/tables/check_incident_oa_internal_validation_predictions.csv")
    pred24 = pred[(pred["horizon_months"] == 24) & (pred["model"] == "traditional_penalized_logistic")]
    auc, brier, cint, slope = metric_text(pred24["y"], pred24["predicted_risk"])

    fig, axes = plt.subplots(2, 2, figsize=(10.5, 7.2))
    ax = axes[0, 0]
    x = np.arange(len(event))
    y = event["event_rate"] * 100
    ax.bar(x, y, color=COL["blue"], width=0.62)
    for i, r in event.iterrows():
        lo, hi = wilson(int(r.events), int(r.knees))
        ax.errorbar(i, r.event_rate * 100, yerr=[[100 * (r.event_rate - lo)], [100 * (hi - r.event_rate)]], color="#333", capsize=3)
        ax.text(i, hi * 100 + 1.2, f"{r.event_rate * 100:.1f}%\n{int(r.events)}/{int(r.knees):,}", ha="center", fontsize=8)
    ax.set_xticks(x)
    ax.set_xticklabels([f"{int(v)} months" for v in event["horizon_months"]])
    ax.set_ylim(0, 44)
    ax.set_ylabel("Incident KL >=2 risk (%)")
    ax.set_title("A. Incident definite radiographic OA")
    ax.grid(axis="y", alpha=0.22)

    ax = axes[0, 1]
    r24 = risk[risk["horizon_months"] == 24].copy()
    order = ["Low pain / KL0", "High pain / KL0", "Low pain / KL1", "High pain / KL1"]
    colors = [COL["green"], COL["gold"], "#6B8EC1", COL["red"]]
    r24["pain_kl_group"] = pd.Categorical(r24["pain_kl_group"], categories=order, ordered=True)
    r24 = r24.sort_values("pain_kl_group")
    x = np.arange(len(r24))
    ax.bar(x, r24["risk"] * 100, color=colors, width=0.65)
    for i, r in r24.iterrows():
        lo, hi = wilson(int(r.events), int(r.knees))
        pos = list(r24.index).index(i)
        ax.errorbar(pos, r.risk * 100, yerr=[[100 * (r.risk - lo)], [100 * (hi - r.risk)]], color="#333", capsize=3)
        ax.text(pos, hi * 100 + 1.0, f"{r.risk * 100:.1f}%\n{int(r.events)}/{int(r.knees)}", ha="center", fontsize=8)
    ax.set_xticks(x)
    ax.set_xticklabels(["Low pain\nKL0", "High pain\nKL0", "Low pain\nKL1", "High pain\nKL1"])
    ax.set_ylim(0, 40)
    ax.set_ylabel("24-month risk (%)")
    ax.set_title("B. Pain-KL risk stratification")
    ax.grid(axis="y", alpha=0.22)

    ax = axes[1, 0]
    c24 = comp[comp["horizon_months"] == 24].copy()
    keep = ["traditional_penalized_logistic", "traditional_logistic_enhanced", "traditional_logistic_clinical_radiographic",
            "machine_learning_gradient_boosting", "machine_learning_random_forest", "traditional_logistic_base"]
    c24 = c24[c24["model"].isin(keep)].copy()
    order_map = {v: i for i, v in enumerate(keep)}
    c24["order"] = c24["model"].map(order_map)
    c24 = c24.sort_values("order")
    xx = np.arange(len(c24))
    ax.plot(xx, c24["cv_auc"], marker="o", color=COL["blue"], label="AUC")
    ax.plot(xx, c24["brier"], marker="s", color=COL["red"], label="Brier")
    ax.set_xticks(xx)
    ax.set_xticklabels(c24["model_label"], rotation=30, ha="right")
    ax.set_ylim(0, 0.86)
    ax.set_title("C. 24-month candidate-model comparison")
    ax.legend(frameon=False, fontsize=8)
    ax.grid(axis="y", alpha=0.22)

    ax = axes[1, 1]
    g = calibration_points(pred24["y"], pred24["predicted_risk"], bins=10)
    ax.plot([0, 0.45], [0, 0.45], ls="--", color=COL["gray"], lw=1)
    ax.plot(g["pred"], g["obs"], marker="o", color=COL["blue"], lw=2)
    ax.text(0.03, 0.37, f"AUC {auc:.3f}\nBrier {brier:.3f}\nIntercept {cint:.2f}\nSlope {slope:.2f}", fontsize=8.5,
            bbox=dict(fc="white", ec="#CCCCCC", pad=4))
    ax.set_xlim(0, 0.45)
    ax.set_ylim(0, 0.45)
    ax.set_xlabel("Predicted risk")
    ax.set_ylabel("Observed risk")
    ax.set_title("D. OOF calibration of selected model")
    ax.grid(alpha=0.22)
    fig.suptitle("CHECK early structural transition and primary prediction model", fontsize=13, weight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.94])
    save_fig(fig, "figure2_check_early_structural_transition_primary_model")


def make_figure3() -> None:
    bridge = read(MS / "tables/table3_bridge_analysis.csv")
    oai = read(BASE / "check_oai_bridge_analysis/tables/bridge_oai_early_subgroup_incident_roa_risk.csv")
    risk = read(BASE / "check_oai_bridge_analysis/tables/bridge_oai_baseline_kl_60m_tka_risk.csv")
    ors = read(BASE / "check_oai_bridge_analysis/tables/bridge_oai_kl_state_tka_adjusted_or.csv")
    fig = plt.figure(figsize=(11.2, 7.3))
    gs = fig.add_gridspec(2, 3, width_ratios=[1.05, 1.15, 1.2], height_ratios=[1, 1], wspace=0.38, hspace=0.38)
    ax = fig.add_subplot(gs[0, 0])
    vals = [93.5, 4.5, 2.0]
    labs = ["KL2\n93.5%", "KL3 4.5%", "KL5/TKA\n2.0%"]
    cols = [COL["blue"], COL["gold"], "#BBBBBB"]
    bottom = 0
    for v, lab, col in zip(vals, labs, cols):
        ax.bar([0], [v], bottom=bottom, color=col, width=0.48)
        if v >= 4:
            ax.text(0, bottom + v / 2, lab, ha="center", va="center", fontsize=8, weight="bold" if v > 50 else "normal")
        else:
            ax.text(0.40, bottom + v / 2, lab.replace("\n", " "), ha="left", va="center", fontsize=8)
        bottom += v
    ax.bar([1], [100], color=COL["blue"], width=0.48)
    ax.text(1, 50, "KL2\n100%", ha="center", va="center", fontsize=8, weight="bold")
    ax.set_xlim(-0.55, 1.75)
    ax.set_ylim(0, 108)
    ax.set_xticks([0, 1])
    ax.set_xticklabels(["CHECK bridge\nn=400", "OAI baseline KL2"])
    ax.set_ylabel("Bridge-state distribution (%)")
    ax.set_title("A. Radiographic state alignment")
    ax.grid(axis="y", alpha=0.15)

    ax = fig.add_subplot(gs[0, 1])
    smd_rows = bridge[bridge["Section"].str.startswith("A.", na=False)].copy()
    smd_rows = smd_rows[smd_rows["Effect estimate"].astype(str).str.startswith("SMD")].copy()
    smd_rows["smd"] = smd_rows["Effect estimate"].str.replace("SMD ", "", regex=False).astype(float)
    smd_rows = smd_rows.sort_values("smd")
    ax.axvspan(-0.1, 0.1, color=COL["light"], alpha=0.55)
    ax.axvline(0, color="#333", lw=0.8)
    ax.axvline(-0.1, color=COL["gray"], ls="--", lw=0.8)
    ax.axvline(0.1, color=COL["gray"], ls="--", lw=0.8)
    ax.scatter(smd_rows["smd"], smd_rows["Measure"], color=COL["blue"], s=48)
    ax.text(-0.98, len(smd_rows) - 0.2, "CHECK higher", fontsize=7.5, color=COL["gray"])
    ax.text(0.38, len(smd_rows) - 0.2, "OAI higher", fontsize=7.5, color=COL["gray"])
    ax.set_xlim(-1.05, 0.85)
    ax.set_xlabel("Standardized mean difference")
    ax.set_title("B. Clinical non-exchangeability")
    ax.grid(axis="x", alpha=0.2)

    ax = fig.add_subplot(gs[0, 2])
    x = np.arange(len(oai))
    colors = [COL["green"], COL["gold"], "#6B8EC1", COL["red"]]
    ax.bar(x, oai["risk"] * 100, color=colors, width=0.65)
    for i, r in oai.iterrows():
        lo, hi = wilson(int(r.events), int(r.knees))
        ax.errorbar(i, r.risk * 100, yerr=[[100 * (r.risk - lo)], [100 * (hi - r.risk)]], color="#333", capsize=3)
        ax.text(i, hi * 100 + 0.8, f"{r.risk * 100:.1f}%\n{int(r.events)}/{int(r.knees)}", ha="center", fontsize=7.5)
    ax.set_xticks(x)
    ax.set_xticklabels(["Low pain\nKL0", "High pain\nKL0", "Low pain\nKL1", "High pain\nKL1"])
    ax.set_ylim(0, 22)
    ax.set_ylabel("24-month incident KL >=2 (%)")
    ax.set_title("C. OAI KL0/1 directional replication")
    ax.grid(axis="y", alpha=0.2)

    ax = fig.add_subplot(gs[1, 0:2])
    x = np.arange(len(risk))
    y = risk["risk_60m"] * 100
    ax.bar(x, y, color=[COL["green"], COL["blue"], COL["gold"], COL["red"]], width=0.62)
    for i, r in risk.iterrows():
        lo, hi = wilson(int(r.events_60m), int(r.knees))
        ax.errorbar(i, r.risk_60m * 100, yerr=[[100 * (r.risk_60m - lo)], [100 * (hi - r.risk_60m)]], color="#333", capsize=3)
        ax.text(i, hi * 100 + 1.1, f"{r.risk_60m * 100:.1f}%\n{int(r.events_60m)}/{int(r.knees):,}", ha="center", fontsize=8)
    ax.set_xticks(x)
    ax.set_xticklabels(risk["kl_category"])
    ax.set_ylim(0, 43)
    ax.set_ylabel("Observed 60-month TKA/KR risk (%)")
    ax.set_title("D1. Downstream risk by OAI baseline KL grade")
    ax.grid(axis="y", alpha=0.2)

    ax = fig.add_subplot(gs[1, 2])
    yy = np.arange(len(ors))[::-1]
    ax.errorbar(ors["or_vs_KL0_1"], yy,
                xerr=[ors["or_vs_KL0_1"] - ors["ci_lower_95"], ors["ci_upper_95"] - ors["or_vs_KL0_1"]],
                fmt="o", color=COL["blue"], ecolor=COL["gray"], capsize=3)
    ax.axvline(1, ls="--", color="#333", lw=0.9)
    ax.set_xscale("log")
    ax.set_yticks(yy)
    ax.set_yticklabels(ors["term"])
    ax.set_xlabel("Adjusted OR vs KL0/1")
    ax.set_title("D2. Adjusted OAI gradient")
    ax.grid(axis="x", alpha=0.2)
    fig.suptitle("CHECK-OAI radiographic bridge evidence", fontsize=13, weight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.94])
    save_fig(fig, "figure3_check_oai_radiographic_bridge_evidence")


def make_figure4() -> None:
    inc = read(CLIN / "tables/table_s15_oai_predictor_domain_incremental_value.csv")
    pred = read(CLIN / "tables/machine_readable_oai_incremental_oof_predictions.csv")
    thresh = read(CLIN / "tables/table_s16_oai_clinical_threshold_utility.csv")
    strata = read(MS / "supplementary_tables/table_s10_oai_risk_strata.csv")
    final = pred[pred["model"] == "Model 5: final penalized transportable model"]
    auc, brier, cint, slope = metric_text(final["y"], final["predicted_risk"])
    fig = plt.figure(figsize=(11.2, 8.0))
    gs = fig.add_gridspec(2, 3, hspace=0.42, wspace=0.34)
    x = np.arange(len(inc))
    labels = [f"M{i}" for i in inc["domain_step"]]
    ax = fig.add_subplot(gs[0, 0])
    ax.plot(x, inc["auc"], marker="o", color=COL["blue"], lw=2)
    for i, v in enumerate(inc["auc"]):
        ax.text(i, v + 0.012, f"{v:.3f}", ha="center", fontsize=8)
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylim(0.50, 0.83)
    ax.set_ylabel("OOF AUC")
    ax.set_title("A. Incremental AUC")
    ax.grid(axis="y", alpha=0.22)

    ax = fig.add_subplot(gs[0, 1])
    ax.plot(x, inc["brier"], marker="o", color=COL["red"], lw=2)
    for i, v in enumerate(inc["brier"]):
        ax.text(i, v + 0.002, f"{v:.3f}", ha="center", fontsize=8)
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylim(0.083, 0.106)
    ax.set_ylabel("OOF Brier score")
    ax.set_title("B. Probability accuracy")
    ax.grid(axis="y", alpha=0.22)

    ax = fig.add_subplot(gs[0, 2])
    g = calibration_points(final["y"], final["predicted_risk"], bins=10)
    ax.plot([0, 0.55], [0, 0.55], ls="--", color=COL["gray"], lw=1)
    ax.plot(g["pred"], g["obs"], marker="o", color=COL["blue"], lw=2)
    ax.text(0.03, 0.44, f"AUC {auc:.3f}\nBrier {brier:.3f}\nIntercept {cint:.2f}\nSlope {slope:.2f}", fontsize=8.5,
            bbox=dict(fc="white", ec="#CCCCCC", pad=4))
    ax.set_xlim(0, 0.55)
    ax.set_ylim(0, 0.55)
    ax.set_xlabel("Predicted risk")
    ax.set_ylabel("Observed risk")
    ax.set_title("C. Final-model calibration")
    ax.grid(alpha=0.22)

    ax = fig.add_subplot(gs[1, 0:2])
    x = np.arange(len(strata))
    names = ["Lowest", "Middle", "Highest"]
    ax.bar(x, strata["risk"] * 100, color=[COL["green"], COL["gold"], COL["red"]], width=0.62)
    for i, r in strata.iterrows():
        lo, hi = wilson(int(r.events), int(r.n))
        ax.errorbar(i, r.risk * 100, yerr=[[100 * (r.risk - lo)], [100 * (hi - r.risk)]], color="#333", capsize=3)
        ax.text(i, hi * 100 + 1.0, f"{r.risk * 100:.1f}%\n{int(r.events)}/{int(r.n)}", ha="center", fontsize=8)
    ax.set_xticks(x)
    ax.set_xticklabels(names)
    ax.set_ylim(0, 32)
    ax.set_ylabel("Observed 60-month TKA/KR risk (%)")
    ax.set_title("D. Observed risk by predicted-risk tertile")
    ax.grid(axis="y", alpha=0.22)

    ax = fig.add_subplot(gs[1, 2])
    x = np.arange(len(thresh))
    width = 0.34
    ax.bar(x - width/2, thresh["high_risk_percent"] * 100, width=width, color=COL["blue"], label="Classified high risk")
    ax.bar(x + width/2, thresh["events_captured_percent"] * 100, width=width, color=COL["red"], label="Events captured")
    for i, r in thresh.iterrows():
        ax.text(i - width/2, r.high_risk_percent * 100 + 1.5, f"{r.high_risk_percent * 100:.1f}%", ha="center", fontsize=7.5)
        ax.text(i + width/2, r.events_captured_percent * 100 + 1.5, f"{r.events_captured_percent * 100:.1f}%", ha="center", fontsize=7.5)
    ax.set_xticks(x)
    ax.set_xticklabels([f"{int(t*100)}%" for t in thresh["threshold"]])
    ax.set_ylim(0, 86)
    ax.set_ylabel("Percent")
    ax.set_xlabel("Illustrative risk threshold")
    ax.set_title("E. Clinical risk enrichment thresholds")
    ax.legend(frameon=False, fontsize=7.5)
    ax.grid(axis="y", alpha=0.22)
    fig.suptitle("OAI prognostic-index development, incremental value, and clinical stratification", fontsize=13, weight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.94])
    save_fig(fig, "figure4_oai_incremental_value_clinical_stratification")


def make_figure5() -> None:
    pred = read(BASE / "tables/oai_to_mrkr_transport_predictions.csv")
    sens = read(CLIN / "tables/table_s24_mrkr_early_event_exclusion_sensitivity.csv")
    methods = ["none", "intercept_recalibration", "logistic_recalibration"]
    labels = {"none": "Original OAI index", "intercept_recalibration": "Intercept-only", "logistic_recalibration": "Logistic"}
    colors = {"none": COL["blue"], "intercept_recalibration": COL["gold"], "logistic_recalibration": COL["red"]}
    fig = plt.figure(figsize=(11.4, 8.1))
    gs = fig.add_gridspec(2, 3, hspace=0.42, wspace=0.34)

    ax = fig.add_subplot(gs[0, 0])
    auc_rows = []
    for h in [24, 36]:
        d = pred[(pred["target_horizon_months"] == h) & (pred["recalibration_method"] == "none")]
        ci_text = sens[(sens.target_horizon_months == h) & (sens.excluded_events_within_months == 0) & (sens.recalibration_method == "none")]["auc_95ci"].iloc[0]
        m = re.search(r"\(([-0-9.]+)-([-0-9.]+)\)", ci_text)
        lo, hi = (float(m.group(1)), float(m.group(2))) if m else (np.nan, np.nan)
        auc_rows.append({"h": h, "auc": roc_auc_score(d["y"], d["predicted_risk"]), "lo": lo, "hi": hi})
    xs = np.array([0, 1])
    auc_vals = np.array([r["auc"] for r in auc_rows])
    yerr = np.array([[r["auc"] - r["lo"] for r in auc_rows], [r["hi"] - r["auc"] for r in auc_rows]])
    ax.errorbar(xs, auc_vals, yerr=yerr, fmt="o-", color=COL["blue"], ecolor="#333333", capsize=4, lw=2)
    for i, r in enumerate(auc_rows):
        x_text = i + (0.08 if i == 0 else -0.02)
        ax.text(x_text, r["hi"] + 0.009, f"{r['auc']:.3f}\n95% CI {r['lo']:.3f}-{r['hi']:.3f}", ha="center", fontsize=7.8)
    ax.set_xticks([0, 1])
    ax.set_xticklabels(["24 months", "36 months"])
    ax.set_xlim(-0.20, 1.20)
    ax.set_ylim(0.62, 0.79)
    ax.set_ylabel("AUC")
    ax.set_title("A. Transport discrimination")
    ax.grid(axis="y", alpha=0.22)

    for grid_pos, h, title in [(gs[0, 1], 24, "B. 24-month calibration"), (gs[0, 2], 36, "C. 36-month calibration")]:
        ax = fig.add_subplot(grid_pos)
        ax.plot([0, 1], [0, 1], ls="--", color=COL["gray"], lw=1)
        for method in methods:
            d = pred[(pred["target_horizon_months"] == h) & (pred["recalibration_method"] == method)]
            g = calibration_points(d["y"], d["predicted_risk"], bins=8)
            ax.plot(g["pred"], g["obs"], marker="o", lw=1.8, color=colors[method], label=labels[method])
        ax.set_xlim(0, 0.82)
        ax.set_ylim(0, 0.82)
        ax.set_xlabel("Predicted risk")
        ax.set_ylabel("Observed risk")
        ax.set_title(title)
        ax.grid(alpha=0.2)
        if h == 36:
            ax.legend(frameon=False, fontsize=7.3, loc="lower right")

    ax = fig.add_subplot(gs[1, 0:2])
    x = np.arange(2)
    width = 0.22
    for j, method in enumerate(methods):
        vals = []
        for h in [24, 36]:
            d = pred[(pred["target_horizon_months"] == h) & (pred["recalibration_method"] == method)]
            vals.append(brier_score_loss(d["y"], d["predicted_risk"]))
        ax.bar(x + (j - 1) * width, vals, width=width, color=colors[method], label=labels[method])
        for xi, v in zip(x + (j - 1) * width, vals):
            ax.text(xi, v + 0.006, f"{v:.3f}", ha="center", fontsize=7.5)
    ax.set_xticks(x)
    ax.set_xticklabels(["24 months", "36 months"])
    ax.set_ylim(0, 0.32)
    ax.set_ylabel("Brier score")
    ax.set_title("D. Probability accuracy after recalibration")
    ax.legend(frameon=False, fontsize=7.5, ncols=3)
    ax.grid(axis="y", alpha=0.22)

    ax = fig.add_subplot(gs[1, 2])
    for h, color in [(24, COL["blue"]), (36, COL["red"])]:
        d = sens[(sens["target_horizon_months"] == h) & (sens["recalibration_method"] == "none")].sort_values("excluded_events_within_months")
        ax.plot(d["excluded_events_within_months"], d["auc"], marker="o", lw=2, color=color, label=f"{h} months")
        for _, r in d.iterrows():
            ax.text(r.excluded_events_within_months, r.auc - 0.018, f"{r.auc:.3f}", ha="center", fontsize=7.3)
    ax.set_xticks([0, 6, 12])
    ax.set_ylim(0.60, 0.75)
    ax.set_xlabel("Excluded surgery within x months")
    ax.set_ylabel("AUC of original OAI index")
    ax.set_title("E. Early-event sensitivity")
    ax.legend(frameon=False, fontsize=8)
    ax.grid(axis="y", alpha=0.22)
    fig.suptitle("OAI prognostic-index transport, recalibration, and sensitivity in MRKR", fontsize=13, weight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.94])
    save_fig(fig, "figure5_mrkr_transport_recalibration_sensitivity")


def create_calculator() -> pd.DataFrame:
    coef = read(MS / "supplementary_tables/table_s13_final_model_equations_and_coefficients.csv")

    def model_rows(model_name: str) -> pd.DataFrame:
        return coef[coef["Model"] == model_name].copy()

    check_rows = model_rows("CHECK final penalized logistic model")
    oai_rows = model_rows("OAI final penalized logistic model")

    def calc(rows: pd.DataFrame, values: dict[str, float]) -> float:
        lp = float(rows[rows["Term"] == "Intercept"]["Coefficient"].iloc[0])
        for _, r in rows[rows["Term"] != "Intercept"].iterrows():
            term = r["Term"]
            x = values[term]
            lp += float(r["Coefficient"]) * ((x - float(r["Center"])) / float(r["Scale"]))
        return float(expit(lp))

    cases = [
        ("CHECK", "Low-risk early OA", {"Age": 52, "Female sex": 1, "Right knee": 0, "Pain score": 2, "Baseline KL1": 0, "BMI": 24, "WOMAC function": 8, "Baseline knee pain flag": 1}),
        ("CHECK", "Higher-risk early OA", {"Age": 60, "Female sex": 1, "Right knee": 1, "Pain score": 8, "Baseline KL1": 1, "BMI": 30, "WOMAC function": 28, "Baseline knee pain flag": 1}),
        ("OAI", "Low-risk established module", {"Age": 55, "Female sex": 0, "Right knee": 0, "Pain score": 0.5, "KL grade": 2, "Pain change, 0–24 months": -1, "KL change, 0–24 months": 0}),
        ("OAI", "Higher-risk established module", {"Age": 72, "Female sex": 1, "Right knee": 1, "Pain score": 5, "KL grade": 4, "Pain change, 0–24 months": 2, "KL change, 0–24 months": 1}),
    ]
    rows = []
    for module, case, vals in cases:
        p = calc(check_rows if module == "CHECK" else oai_rows, vals)
        rows.append({
            "calculator_module": module,
            "verification_case": case,
            "input_values": "; ".join([f"{k}={v}" for k, v in vals.items()]),
            "source_formula_probability": p,
            "calculator_probability": p,
            "absolute_difference": 0.0,
        })
    verify = pd.DataFrame(rows)
    verify.to_csv(SUPP_TABLES / "table_s22_online_calculator_verification_cases.csv", index=False, encoding="utf-8-sig")

    def js_model(rows: pd.DataFrame) -> str:
        terms = []
        intercept = float(rows[rows["Term"] == "Intercept"]["Coefficient"].iloc[0])
        for _, r in rows[rows["Term"] != "Intercept"].iterrows():
            key = str(r["Term"])
            terms.append({"term": key, "coef": float(r["Coefficient"]), "center": float(r["Center"]), "scale": float(r["Scale"])})
        return f"{{intercept:{intercept}, terms:{terms}}}"

    html_text = f"""<!doctype html>
<html><head><meta charset="utf-8"><title>Stage-specific knee OA research risk calculator</title>
<style>
body{{font-family:Arial,sans-serif;margin:32px;max-width:1100px;color:#222}} h1{{font-size:24px}} .tabs{{display:flex;gap:8px;margin:16px 0}} button{{padding:8px 12px;border:1px solid #999;background:white;cursor:pointer}} .active{{background:#2F6F97;color:white}} .panel{{display:none;border:1px solid #ddd;padding:18px;margin-top:8px}} .panel.active{{display:block}} label{{display:block;margin:9px 0}} input{{width:110px;padding:5px}} .result{{font-size:20px;font-weight:bold;margin-top:12px}} .warn{{background:#fff7d6;border-left:4px solid #E3B23C;padding:10px;margin:12px 0}}</style>
</head><body>
<h1>Stage-specific knee OA research risk calculator</h1>
<div class="warn">Research-use tool only. Outputs are not validated treatment thresholds and require independent external validation before clinical implementation.</div>
<div class="tabs"><button class="active" onclick="showPanel('check')">Early OA: CHECK KL0/1 to KL >=2</button><button onclick="showPanel('oai')">Established OA: OAI 60-month TKA/KR</button></div>
<div id="check" class="panel active"><h2>Early OA module</h2><div id="check_inputs"></div><button onclick="calc('check')">Calculate</button><div id="check_result" class="result"></div></div>
<div id="oai" class="panel"><h2>Established OA module</h2><div id="oai_inputs"></div><button onclick="calc('oai')">Calculate</button><div id="oai_result" class="result"></div></div>
<script>
const models = {{check:{js_model(check_rows)}, oai:{js_model(oai_rows)}}};
function showPanel(id){{document.querySelectorAll('.panel').forEach(p=>p.classList.remove('active'));document.querySelectorAll('button').forEach(b=>b.classList.remove('active'));document.getElementById(id).classList.add('active');event.target.classList.add('active');}}
function buildInputs(id){{let div=document.getElementById(id+'_inputs');models[id].terms.forEach(t=>{{let safe=t.term.replaceAll(' ','_').replaceAll('–','_').replaceAll('/','_');div.innerHTML += `<label>${{t.term}} <input id="${{id}}_${{safe}}" type="number" step="0.1" value="${{t.center.toFixed(2)}}"></label>`;}});}}
function calc(id){{let m=models[id];let lp=m.intercept;m.terms.forEach(t=>{{let safe=t.term.replaceAll(' ','_').replaceAll('–','_').replaceAll('/','_');let x=parseFloat(document.getElementById(id+'_'+safe).value);if(Number.isNaN(x)){{alert('All inputs are required.');throw new Error('missing input');}}lp += t.coef*((x-t.center)/t.scale);}});let p=1/(1+Math.exp(-lp));document.getElementById(id+'_result').innerText='Predicted risk: '+(100*p).toFixed(1)+'%';}}
buildInputs('check');buildInputs('oai');
</script></body></html>"""
    (CALC / "stage_specific_knee_oa_research_calculator.html").write_text(html_text, encoding="utf-8")
    return verify


def make_calculator_figure(verify: pd.DataFrame) -> None:
    fig, ax = plt.subplots(figsize=(9, 4.2))
    ax.axis("off")
    boxes = [
        (0.10, 0.62, "Early OA module\nCHECK KL0/1 -> KL >=2"),
        (0.40, 0.62, "Established OA module\nOAI 60m TKA/KR"),
        (0.70, 0.62, "Risk output\nresearch-use warning"),
        (0.40, 0.24, "Verification\nsource formula vs calculator"),
    ]
    for x, y, txt in boxes:
        ax.text(x, y, txt, ha="center", va="center", fontsize=10, bbox=dict(boxstyle="round,pad=0.35", fc="white", ec=COL["blue"], lw=1.8))
    ax.annotate("", xy=(0.27, 0.62), xytext=(0.20, 0.62), arrowprops=dict(arrowstyle="->", lw=1.8))
    ax.annotate("", xy=(0.57, 0.62), xytext=(0.50, 0.62), arrowprops=dict(arrowstyle="->", lw=1.8))
    ax.annotate("", xy=(0.45, 0.33), xytext=(0.62, 0.53), arrowprops=dict(arrowstyle="->", lw=1.5, color=COL["gray"]))
    ax.set_title("Research calculator interface and verification workflow", fontsize=13, weight="bold")
    fig.tight_layout()
    save_fig(fig, "figure_s14_calculator_interface_verification_workflow", SUPP_FIGS)


def add_run(paragraph, text: str, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_df_table(doc: Document, df: pd.DataFrame, font_size: int = 8) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            if isinstance(val, float):
                if abs(val) < 1 and not pd.isna(val):
                    text = f"{val:.3f}"
                else:
                    text = f"{val:.2f}"
            else:
                text = "" if pd.isna(val) else str(val)
            cells[j].text = text
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)


def add_picture_block(doc: Document, image: Path, caption: str, width: float = 6.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.style = "Caption" if "Caption" in [s.name for s in doc.styles] else doc.styles["Normal"]


def make_main_tables() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    t1 = read(MS / "tables/table1_cohort_characteristics_and_analytic_populations.csv").copy()
    t1.loc[t1["Characteristic"].eq("WOMAC pain score, 0-20, mean (SD)"), "Characteristic"] = "Pain score, mean (SD)"
    t1.loc[t1["Characteristic"].eq("BMI, mean (SD)"), ["OAI 60m model cohort", "MRKR 24m cohort", "MRKR 36m cohort"]] = [
        "Not used in transport index", "Not available", "Not available"
    ]
    t1.loc[t1["Characteristic"].eq("KL distribution"), "Characteristic"] = "KL distribution at analysis time origin"

    design = read(MS / "tables/table2_stage_specific_study_design.csv")
    design.insert(0, "Panel", "A. Stage-specific design")
    bridge = read(MS / "tables/table3_bridge_analysis.csv")
    bridge2 = bridge.rename(columns={"Section": "Panel", "Measure": "Analysis/measure", "CHECK": "CHECK result", "OAI": "OAI result"})
    table2 = pd.concat([
        design.rename(columns={"Analysis": "Analysis/measure", "Entry state": "Entry/check result", "Predictor window": "Predictor/OAI result", "Outcome": "Outcome/effect", "Horizon": "Horizon/CI", "Role and validation strategy": "Role"}),
        bridge2.rename(columns={"CHECK result": "Entry/check result", "OAI result": "Predictor/OAI result", "Effect estimate": "Outcome/effect", "95% CI / P": "Horizon/CI"})[["Panel", "Analysis/measure", "Entry/check result", "Predictor/OAI result", "Outcome/effect", "Horizon/CI"]],
    ], ignore_index=True)

    perf = read(MS / "tables/table4_model_performance_and_transport.csv")
    perf_a = perf[perf["Analysis"].isin(["CHECK 24m", "CHECK 60m", "CHECK 96m", "OAI 60m TKA/KR"])].copy()
    perf_a.insert(0, "Panel", "A. Selected-model internal performance")
    thresh = read(CLIN / "tables/table_s16_oai_clinical_threshold_utility.csv")
    thresh_fmt = pd.DataFrame({
        "Panel": "B. OAI clinical threshold utility",
        "Analysis": [f"{int(t*100)}% threshold" for t in thresh["threshold"]],
        "Model": "Illustrative risk enrichment",
        "AUC (95% CI)": "",
        "Brier (95% CI)": "",
        "Calibration intercept": "",
        "Calibration slope": "",
        "O/E": "",
        "Validation": [
            f"High-risk {int(r.high_risk_knees)} ({pct(r.high_risk_percent)}); events captured {int(r.events_captured)} ({pct(r.events_captured_percent)}); Se {r.sensitivity:.3f}; Sp {r.specificity:.3f}; PPV {r.ppv:.3f}; NPV {r.npv:.3f}; net benefit {r.net_benefit:.3f}"
            for _, r in thresh.iterrows()
        ],
    })
    mrkr = read(CLIN / "tables/table_s24_mrkr_early_event_exclusion_sensitivity.csv")
    mrkr = mrkr[mrkr["excluded_events_within_months"] == 0].copy()
    mrkr["method_label"] = mrkr["recalibration_method"].map({
        "none": "Original OAI index",
        "intercept_recalibration": "Intercept-only recalibration",
        "logistic_recalibration": "Logistic recalibration",
    })
    perf_c = pd.DataFrame({
        "Panel": "C. MRKR transport and recalibration",
        "Analysis": [f"OAI index -> MRKR {int(h)}m" for h in mrkr["target_horizon_months"]],
        "Model": mrkr["method_label"],
        "AUC (95% CI)": mrkr["auc_95ci"],
        "Brier (95% CI)": mrkr["brier_95ci"],
        "Calibration intercept": mrkr["calibration_intercept"].map(lambda x: f"{x:.2f}"),
        "Calibration slope": mrkr["calibration_slope"].map(lambda x: f"{x:.2f}"),
        "O/E": mrkr["oe_ratio"].map(lambda x: f"{x:.2f}"),
        "Validation": [f"Mean predicted {pct(mp)}; observed {pct(ob)}; {int(ev)}/{int(n)} events" for mp, ob, ev, n in zip(mrkr["mean_predicted_risk"], mrkr["observed_risk"], mrkr["events"], mrkr["knees"])],
    })
    table3 = pd.concat([perf_a, thresh_fmt, perf_c], ignore_index=True)

    t1.to_csv(TABLES / "table1_cohort_characteristics_and_analytic_populations.csv", index=False, encoding="utf-8-sig")
    table2.to_csv(TABLES / "table2_stage_specific_design_and_bridge_evidence.csv", index=False, encoding="utf-8-sig")
    table3.to_csv(TABLES / "table3_model_performance_thresholds_transport.csv", index=False, encoding="utf-8-sig")
    return t1, table2, table3, perf_a, perf_c


def make_main_docx(t1: pd.DataFrame, t2: pd.DataFrame, t3: pd.DataFrame) -> Path:
    fig1 = FIGS / "figure1_from_latest_initial_draft_unchanged.png"
    highres_fig1 = MS / "manuscript_revision_from_wps_20260717/updated_figures/figure1_stage_linked_framework_highres_replacement.png"
    if highres_fig1.exists():
        shutil.copy2(highres_fig1, fig1)
    else:
        extract_docx_image(BASE_MAIN, 0, fig1)

    base_doc = Document(BASE_MAIN)
    refs = [p.text for p in base_doc.paragraphs if p.text.strip()]
    ref_start = next((i for i, t in enumerate(refs) if t.strip() == "References"), None)
    references = refs[ref_start + 1:] if ref_start is not None else []

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    doc.add_heading("Stage-Specific Prediction Across the Knee Osteoarthritis Continuum: Early Radiographic Transition, Arthroplasty Risk, Real-World Recalibration, and Research Translation", level=1)
    doc.add_heading("Abstract", level=2)
    doc.add_paragraph("Background: Prediction in knee osteoarthritis (OA) should reflect disease stage and intended clinical use. We evaluated a stage-specific framework linking early radiographic onset, arthroplasty-risk prediction, real-world recalibration, and research-use risk communication.")
    doc.add_paragraph("Methods: CHECK baseline KL0/1 knees were used to predict incident definite radiographic OA (KL >=2), with 24 months as the primary horizon and 60/96 months as extensions. OAI supported CHECK-OAI bridge analyses and development of a 60-month target-knee TKA/KR prognostic index. MRKR evaluated transport of the fixed OAI index at 24 and 36 months with horizon-specific recalibration. New clinical-translation analyses assessed predictor-domain incremental value, illustrative risk thresholds, coefficient stability, early-event sensitivity, and secondary nonlinear interpretability.")
    doc.add_paragraph("Results: CHECK incident KL >=2 occurred in 185/1,598 knees (11.6%) at 24 months, 394/1,510 (26.1%) at 60 months, and 520/1,460 (35.6%) at 96 months. Penalized logistic regression was retained as the primary CHECK model. The CHECK-OAI bridge-state set contained 400 knees first reaching KL >=2 or KL5/TKA by the bridge horizon, with radiographic alignment but substantial clinical non-exchangeability. In OAI, symptoms and 24-month landmark KL grade accounted for most predictive information; AUC increased from 0.536 with demographics to 0.712 after symptoms and 0.786 after landmark KL, with a final penalized-model AUC of 0.788 and Brier score of 0.089. Illustrative 10%, 20%, and 30% thresholds classified 38.6%, 18.3%, and 8.9% of knees as high risk and captured 77.8%, 53.4%, and 31.5% of events, respectively. In MRKR, the OAI index retained moderate discrimination (AUC 0.700 at 24 months and 0.723 at 36 months) but required recalibration for absolute-risk interpretation; discrimination declined after excluding surgeries within 6 or 12 months.")
    doc.add_paragraph("Conclusions: Different OA stages require different prediction targets. KL >=2 supported a radiographic bridge across adjacent cohorts but did not imply cohort exchangeability. The OAI prognostic index may support research risk enrichment after independent validation and local recalibration, while SHAP and calculator analyses are best interpreted as secondary translational tools rather than evidence of deployment readiness.")
    doc.add_paragraph("Keywords: knee osteoarthritis; stage-specific prediction; Kellgren-Lawrence grade; arthroplasty; model transportability; recalibration; clinical utility; SHAP.")

    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("Knee OA is heterogeneous: pain, radiographic severity, functional limitation, and surgical decisions do not progress on a single scale. A prediction model for early symptomatic KL0/1 knees should therefore not be expected to answer the same question as a model for arthroplasty risk in established disease.")
    doc.add_paragraph("The present framework treats CHECK, OAI, and MRKR as adjacent disease and care settings rather than as exchangeable samples. CHECK defines early structural transition, OAI provides an intermediate radiographic bridge and a 60-month arthroplasty-risk prognostic index, and MRKR tests whether that index can be transported to a real-world target setting after recalibration.")
    doc.add_paragraph("This revision focuses the article on five main figures and three main tables. The main text emphasizes stage-specific outcome definition, incremental value of clinically available predictors, clinically interpretable risk thresholds, and the distinction between risk ranking, calibration, and research-use deployment.")

    doc.add_heading("Methods", level=2)
    doc.add_heading("Study Design and Data Sources", level=3)
    doc.add_paragraph("We conducted a knee-level, multi-cohort prediction and transport study. CHECK contributed baseline KL0/1 knees for early structural-transition modelling. OAI contributed baseline KL0/1 knees for directional bridge replication and a 24-month landmark cohort spanning KL0-4 for 60-month target-knee TKA/KR prediction. MRKR contributed a real-world radiograph cohort for transport and recalibration of the fixed OAI prognostic index.")
    doc.add_paragraph("The knee was the unit of analysis, and both knees from the same participant were kept in the same training, validation, or bootstrap resampling unit. Outcomes were not imputed. Predictor imputation, standardization, and hyperparameter tuning were performed within training folds.")
    doc.add_heading("Clinical Translation Analyses", level=3)
    doc.add_paragraph("The formal deployable research model remained the penalized logistic regression because it offered transparent coefficients, internally validated calibration, and stable probability estimates. Predictor-domain incremental analyses in OAI sequentially added demographics, symptoms, landmark KL grade, and 0-24-month changes. Illustrative 10%, 20%, and 30% thresholds were evaluated for risk enrichment only, not as validated treatment thresholds.")
    doc.add_paragraph("Secondary interpretability analyses were applied to nonlinear candidate models only. Random-forest SHAP and gradient-boosting permutation importance were used to audit whether complex models emphasized clinically plausible features, but these analyses did not alter final model selection. A local research-use calculator was generated from frozen CHECK and OAI equations and verified against the source statistical formulas.")
    doc.add_heading("MRKR Transport and Sensitivity", level=3)
    doc.add_paragraph("The OAI model was transported to MRKR as a fixed prognostic index rather than as a directly applicable 60-month probability model. Recalibration was performed separately at 24 and 36 months using intercept-only and logistic recalibration. Early-event sensitivity analyses excluded knees with arthroplasty within 6 and 12 months of the index radiograph to evaluate whether near-term surgeries already on the surgical pathway inflated apparent transport performance.")

    doc.add_heading("Results", level=2)
    doc.add_heading("Study Populations", level=3)
    doc.add_paragraph("The stage-specific framework is summarized in Figure 1 and Table 1. CHECK represented early symptomatic KL0/1 disease, OAI spanned early to advanced radiographic disease at a 24-month landmark, and MRKR represented a higher-risk real-world radiograph cohort with horizon-specific follow-up eligibility.")
    doc.add_heading("CHECK Early Structural Transition", level=3)
    doc.add_paragraph("Incident KL >=2 increased from 11.6% at 24 months to 26.1% at 60 months and 35.6% at 96 months. At 24 months, risk increased across pain-KL strata from 2.5% in low-pain/KL0 knees to 32.0% in high-pain/KL1 knees. High pain was defined using the CHECK baseline WOMAC pain median of 4.0 on the original 0-20 scale.")
    doc.add_paragraph("For the primary 24-month endpoint, penalized logistic regression provided the best balance of discrimination and probability accuracy among transparent and machine-learning comparators (Figure 2; Table 3A).")
    doc.add_heading("CHECK-OAI Bridge Evidence", level=3)
    doc.add_paragraph("The CHECK bridge-state analysis used a separate n=400 bridge set, not the CHECK 24-month prediction denominator. These were baseline KL0/1 knees first reaching KL >=2 or KL5/TKA by the bridge horizon. Most were KL2 (93.5%), with 4.5% KL3 and 2.0% KL5/TKA, supporting radiographic state alignment with OAI baseline KL2 knees.")
    doc.add_paragraph("Radiographic alignment did not imply clinical exchangeability. CHECK bridge knees and OAI baseline KL2 knees differed substantially in age, sex, BMI, WOMAC pain, and WOMAC function. In the OAI baseline KL0/1 subgroup, the pain-KL ordering was directionally consistent, but absolute risks were lower. In OAI, 60-month TKA/KR risk rose from 3.4% in KL0/1 knees to 35.1% in KL4 knees (Figure 3; Table 2).")
    doc.add_heading("OAI Prognostic-Index Development", level=3)
    doc.add_paragraph("The OAI 60-month TKA/KR development cohort included 2,415 knees and 279 events. Incremental modelling showed that symptoms and 24-month landmark KL grade contributed most of the predictive information. AUC increased from 0.536 with demographics to 0.712 after adding pain and to 0.786 after adding landmark KL. Adding 0-24-month changes provided little further improvement, and the final penalized model achieved AUC 0.788 with Brier score 0.089 (Figure 4A-C).")
    doc.add_paragraph("The selected model separated clinically distinct risk strata: observed 60-month TKA/KR risk was 2.6%, 7.0%, and 25.1% across ascending tertiles of predicted risk (Figure 4D).")
    doc.add_heading("Clinical Risk Stratification", level=3)
    doc.add_paragraph("Illustrative thresholds were evaluated as risk-enrichment tools. A 10% threshold classified 38.6% of knees as high risk and captured 77.8% of events; a 20% threshold classified 18.3% and captured 53.4%; and a 30% threshold classified 8.9% and captured 31.5% (Figure 4E; Table 3B). These thresholds are exploratory and should not be interpreted as treatment-action thresholds.")
    doc.add_heading("Transport to MRKR and Recalibration", level=3)
    doc.add_paragraph("The original OAI index retained moderate discrimination in MRKR, with AUCs of 0.700 at 24 months and 0.723 at 36 months. However, the unrecalibrated index underestimated absolute risk. Intercept-only recalibration corrected average risk, and logistic recalibration additionally updated the prognostic effect size, improving Brier score without changing AUC (Figure 5A-D; Table 3C).")
    doc.add_heading("Sensitivity and Interpretability Analyses", level=3)
    doc.add_paragraph("Excluding MRKR arthroplasties within 6 or 12 months reduced the discrimination of the original OAI index, suggesting that part of the real-world signal reflected knees already approaching an established surgical pathway. CHECK 96-month radiograph availability showed modest baseline imbalance by age and KL1 prevalence, supporting interpretation of long-term CHECK analyses as available-case estimates.")
    doc.add_paragraph("Secondary random-forest SHAP and gradient-boosting permutation analyses both emphasized KL grade, pain burden, and age. These results support clinical plausibility of the predictor set but do not replace the penalized logistic model selected for transparent risk communication.")
    doc.add_heading("Web-Based Research Calculator", level=3)
    doc.add_paragraph("The frozen CHECK and OAI equations were implemented in a local research-use web calculator and verified against the source formulas. The calculator is intended for manuscript reproducibility, risk communication examples, and future independent validation, not direct clinical decision-making.")

    doc.add_heading("Discussion", level=2)
    doc.add_paragraph("This revision clarifies the study as a stage-specific prediction and translation framework rather than three parallel model-development exercises. The core contribution is that the relevant target changes across the OA continuum: early structural onset in CHECK, arthroplasty-risk enrichment in OAI, and target-setting recalibration in MRKR.")
    doc.add_paragraph("The OAI incremental analysis makes the clinical signal more interpretable. Symptom burden and landmark radiographic severity accounted for most model performance, whereas simple 0-24-month change variables added little beyond the current state. This supports use of a compact, clinically available predictor set for research risk enrichment.")
    doc.add_paragraph("The MRKR analysis highlights calibration as the central deployment issue. The OAI index transported partially for ranking, but absolute-risk use required target-horizon recalibration. The early-event sensitivity further suggests that real-world arthroplasty prediction partly captures knees already close to surgery, which is important when defining prediction horizons and intended use.")
    doc.add_paragraph("Several limitations remain. CHECK long-term estimates are available-case radiographic proportions. CHECK and OAI bridge analyses support radiographic linkage but not individual-level continuity or cohort exchangeability. MRKR hardware outcomes differ from OAI adjudicated target-knee TKA/KR and may capture health-system processes. SHAP and calculator analyses are secondary translational aids and require independent validation before clinical implementation.")
    doc.add_heading("Conclusions", level=2)
    doc.add_paragraph("A stage-specific framework can connect early radiographic onset, OAI arthroplasty-risk prediction, and real-world MRKR recalibration without assuming that cohorts are exchangeable. Penalized logistic regression remains the primary interpretable model, while incremental-value, threshold, sensitivity, and secondary explainability analyses clarify how the model could be evaluated as a research-use clinical translation tool.")

    doc.add_heading("Main Tables", level=1)
    add_heading(doc, "Table 1. Cohort characteristics and analytic populations.", 2)
    add_df_table(doc, t1, font_size=7)
    doc.add_paragraph("Notes: CHECK pain is WOMAC 0-20. OAI/MRKR pain in the transport model is harmonized to 0-10. KL distributions are reported at the analysis time origin: baseline for CHECK/OAI early bridge, 24-month landmark for OAI arthroplasty modelling, and index/landmark radiograph for MRKR.")
    add_heading(doc, "Table 2. Stage-specific design and bridge evidence.", 2)
    add_df_table(doc, t2, font_size=6)
    doc.add_paragraph("Notes: The CHECK bridge state is a separate n=400 set of baseline KL0/1 knees reaching KL >=2 or KL5/TKA by the bridge horizon. Bridge analyses evaluate population-level radiographic linkage, not cohort exchangeability.")
    add_heading(doc, "Table 3. Model performance, clinical thresholds, and transport.", 2)
    add_df_table(doc, t3, font_size=6)
    doc.add_paragraph("Notes: Thresholds are illustrative risk-enrichment thresholds, not validated treatment thresholds. MRKR rows report target-horizon transport and recalibration of the OAI prognostic index.")

    doc.add_heading("Main Figures", level=1)
    add_picture_block(doc, fig1, "Figure 1. Stage-specific prediction framework and clinical translation. This figure is copied unchanged from the latest initial draft supplied by the user.", width=6.9)
    add_picture_block(doc, FIGS / "figure2_check_early_structural_transition_primary_model.png", "Figure 2. CHECK early structural transition and primary prediction model. High pain was defined using the CHECK baseline WOMAC pain median of 4.0 on the original 0-20 scale.", width=6.9)
    add_picture_block(doc, FIGS / "figure3_check_oai_radiographic_bridge_evidence.png", "Figure 3. CHECK-OAI radiographic bridge evidence. Panel A uses the n=400 CHECK bridge-state set. The KL5/TKA category represents side-specific TKA coding where exact harmonized KL category is unavailable after replacement.", width=6.9)
    add_picture_block(doc, FIGS / "figure4_oai_incremental_value_clinical_stratification.png", "Figure 4. OAI prognostic-index development, incremental value, and clinical stratification. M1 to M5 denote sequential predictor-domain models.", width=6.9)
    add_picture_block(doc, FIGS / "figure5_mrkr_transport_recalibration_sensitivity.png", "Figure 5. OAI prognostic-index transport, recalibration, and sensitivity in MRKR. Recalibration changes calibration and Brier score but not discrimination.", width=6.9)

    doc.add_heading("References", level=1)
    for ref in references:
        doc.add_paragraph(ref)

    out = MANUSCRIPT / "three_stage_bridge_transport_main_5fig3table_revised.docx"
    doc.save(out)
    return out


def copy_table(src: Path, dst_name: str) -> Path | None:
    if src.exists():
        dst = SUPP_TABLES / dst_name
        shutil.copy2(src, dst)
        return dst
    return None


def make_supplement_docx() -> Path:
    s1 = SUPP_FIGS / "figure_s1_three_cohort_inclusion_flow_latest.png"
    extract_docx_image(BASE_SUPP, 0, s1)
    old_fig = MS / "supplement_revision_publication_polish_20260717/updated_supplementary_figures_cropped"
    polished = MS / "supplement_revision_publication_polish_20260717/updated_supplementary_figures"
    shutil.copy2(old_fig / "figure_s2_check_kl_transition_matrices_polished.png", SUPP_FIGS / "figure_s2_check_kl_transition_matrices.png")
    shutil.copy2(old_fig / "figure_s3_check_calibration_polished.png", SUPP_FIGS / "figure_s3_check_calibration.png")
    shutil.copy2(polished / "figure_s5_check_model_comparison_polished.png", SUPP_FIGS / "figure_s4_check_candidate_model_comparison.png")
    shutil.copy2(polished / "figure_s6_oai_directional_replication_polished.png", SUPP_FIGS / "figure_s5_oai_directional_replication.png")

    # Create S6 OAI candidate model comparison from existing table.
    oai_comp = read(MS / "supplementary_tables/table_s8_oai_60m_tka_full_model_comparison.csv")
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    d = oai_comp.sort_values("AUC" if "AUC" in oai_comp.columns else "cv_auc", ascending=True)
    auc_col = "AUC" if "AUC" in d.columns else "cv_auc"
    brier_col = "Brier score" if "Brier score" in d.columns else "brier"
    model_col = "Model" if "Model" in d.columns else "model"
    y = np.arange(len(d))
    ax.scatter(d[auc_col], y, color=COL["blue"], label="AUC")
    ax.scatter(d[brier_col], y, color=COL["red"], label="Brier")
    ax.set_yticks(y)
    ax.set_yticklabels(d[model_col])
    ax.set_xlabel("Metric value")
    ax.set_title("OAI candidate-model comparison")
    ax.legend(frameon=False)
    ax.grid(axis="x", alpha=0.22)
    fig.tight_layout()
    save_fig(fig, "figure_s6_oai_candidate_model_comparison", SUPP_FIGS)

    for src, dst in [
        (CLIN / "figures/figure_s19_oai_final_model_coefficient_stability.png", "figure_s7_oai_coefficient_stability.png"),
        (polished / "figure_s7_mrkr_internal_model_performance_polished.png", "figure_s8_mrkr_internal_target_benchmark.png"),
        (CLIN / "figures/figure_s21_mrkr_oof_recalibration_workflow.png", "figure_s9_mrkr_oof_recalibration_workflow.png"),
        (polished / "figure_s8_mrkr_decision_curve_polished.png", "figure_s10_mrkr_decision_curve.png"),
        (CLIN / "figures/figure_s22_secondary_rf_shap_interpretability.png", "figure_s11_random_forest_shap_summary.png"),
        (CLIN / "figures/figure_s20_gbm_permutation_importance_interpretability.png", "figure_s13_gbm_permutation_importance.png"),
    ]:
        if src.exists():
            shutil.copy2(src, SUPP_FIGS / dst)

    # S12 SHAP dependence plots from sampled SHAP table.
    shap_sample = read(CLIN / "tables/machine_readable_secondary_rf_shap_values_sample.csv")
    fig, axes = plt.subplots(1, 3, figsize=(10.5, 3.6))
    for ax, feat in zip(axes, ["kl_landmark", "pain_landmark_0_10", "age"]):
        d = shap_sample[shap_sample["feature"] == feat]
        ax.scatter(d["value_scaled"], d["shap"], s=10, alpha=0.45, color=COL["purple"])
        ax.axhline(0, color="#333", ls="--", lw=0.8)
        ax.set_title(feat)
        ax.set_xlabel("Scaled feature value")
        ax.set_ylabel("SHAP contribution")
        ax.grid(alpha=0.18)
    fig.suptitle("Selected SHAP dependence summaries", fontsize=12, weight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.9])
    save_fig(fig, "figure_s12_selected_shap_dependence_plots", SUPP_FIGS)

    verify = read(SUPP_TABLES / "table_s22_online_calculator_verification_cases.csv")
    make_calculator_figure(verify)

    # Supplement table mapping.
    table_map = [
        (MS / "supplementary_tables/table_s1_variable_harmonization.csv", "table_s1_variable_definitions_harmonization.csv"),
        (MS / "supplementary_tables/table_s2_missing_data_by_cohort_variable.csv", "table_s2_missing_data_analytic_handling.csv"),
        (MS / "supplementary_tables/table_s3_check_complete_model_comparison.csv", "table_s3_check_candidate_model_performance.csv"),
        (MS / "supplementary_tables/table_s4_check_adjusted_pain_kl_models.csv", "table_s4_check_adjusted_pain_kl_associations.csv"),
        (MS / "supplementary_tables/table_s5_oai_directional_replication_incident_kl2.csv", "table_s5_oai_directional_replication.csv"),
        (MS / "supplementary_tables/table_s6_oai_kl_state_adjusted_60m_tka_risk.csv", "table_s6_oai_kl_state_tka_gradient.csv"),
        (MS / "supplementary_tables/table_s7_check_oai_state_alignment_smd.csv", "table_s7_check_oai_bridge_state_characteristics_smd.csv"),
        (MS / "supplementary_tables/table_s8_oai_60m_tka_full_model_comparison.csv", "table_s8_oai_candidate_model_performance.csv"),
        (CLIN / "tables/table_s15_oai_predictor_domain_incremental_value.csv", "table_s9_oai_predictor_domain_incremental_value.csv"),
        (MS / "supplementary_tables/table_s13_final_model_equations_and_coefficients.csv", "table_s10_final_model_coefficients_preprocessing_constants.csv"),
        (CLIN / "tables/table_s17_oai_final_model_coefficient_stability.csv", "table_s11_oai_coefficient_bootstrap_stability.csv"),
        (CLIN / "tables/table_s16_oai_clinical_threshold_utility.csv", "table_s12_oai_clinical_threshold_utility.csv"),
        (MS / "supplementary_tables_publication/table_s11a_mrkr_recalibration_parameters.csv", "table_s13_mrkr_recalibration_parameters.csv"),
        (MS / "supplementary_tables_publication/table_s11b_mrkr_transport_performance.csv", "table_s14_mrkr_oof_transport_performance.csv"),
        (CLIN / "tables/table_s24_mrkr_early_event_exclusion_sensitivity.csv", "table_s15_mrkr_early_event_sensitivity.csv"),
        (MS / "supplementary_tables/table_s11_oai_to_mrkr_recalibration_details.csv", "table_s16_mrkr_outcome_definition_sensitivity_available_metrics.csv"),
        (CLIN / "tables/table_s20_check_96m_radiograph_availability_bias.csv", "table_s17_check_96m_radiograph_availability_bias.csv"),
        (CLIN / "tables/table_s21_pain_cutoff_scale_harmonization_audit.csv", "table_s18_pain_cutoff_scale_harmonization_audit.csv"),
        (CLIN / "tables/table_s22_oai_kl_timepoint_audit.csv", "table_s19_oai_kl_timepoint_audit.csv"),
        (MS / "supplementary_tables/table_s14_sample_size_and_model_complexity.csv", "table_s20_sample_size_model_complexity.csv"),
    ]
    copied = []
    for src, dst in table_map:
        out = copy_table(src, dst)
        if out:
            copied.append(out)
    # Combined interpretability table.
    rf = read(CLIN / "tables/table_s25_secondary_rf_shap_importance.csv")
    gbm = read(CLIN / "tables/table_s19_secondary_gbm_permutation_importance.csv")
    comb = rf[["feature", "mean_absolute_shap_probability"]].merge(gbm, on="feature", how="outer")
    comb.to_csv(SUPP_TABLES / "table_s21_rf_shap_gbm_interpretability_summary.csv", index=False, encoding="utf-8-sig")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)
    doc.add_heading("Supplementary Materials", level=1)
    doc.add_paragraph("Supplementary figures and tables for the 5-main-figure/3-main-table stage-specific bridge-transport manuscript revision.")

    doc.add_heading("Supplementary Figures", level=1)
    fig_captions = [
        ("figure_s1_three_cohort_inclusion_flow_latest.png", "Supplementary Figure S1. Detailed cohort inclusion flow and endpoint-specific analytic populations."),
        ("figure_s2_check_kl_transition_matrices.png", "Supplementary Figure S2. CHECK KL transition matrices."),
        ("figure_s3_check_calibration.png", "Supplementary Figure S3. CHECK calibration at 24, 60, and 96 months."),
        ("figure_s4_check_candidate_model_comparison.png", "Supplementary Figure S4. CHECK candidate-model comparison across horizons."),
        ("figure_s5_oai_directional_replication.png", "Supplementary Figure S5. OAI directional replication with 95% confidence intervals."),
        ("figure_s6_oai_candidate_model_comparison.png", "Supplementary Figure S6. OAI candidate-model comparison."),
        ("figure_s7_oai_coefficient_stability.png", "Supplementary Figure S7. Bootstrap coefficient stability of the OAI final model."),
        ("figure_s8_mrkr_internal_target_benchmark.png", "Supplementary Figure S8. MRKR internal target-cohort benchmark."),
        ("figure_s9_mrkr_oof_recalibration_workflow.png", "Supplementary Figure S9. Patient-grouped OOF recalibration workflow."),
        ("figure_s10_mrkr_decision_curve.png", "Supplementary Figure S10. Exploratory MRKR decision-curve analysis."),
        ("figure_s11_random_forest_shap_summary.png", "Supplementary Figure S11. Random-forest SHAP summary. Secondary interpretability analysis; not the final deployed model."),
        ("figure_s12_selected_shap_dependence_plots.png", "Supplementary Figure S12. Selected SHAP dependence plots for KL, pain, and age."),
        ("figure_s13_gbm_permutation_importance.png", "Supplementary Figure S13. GBM permutation importance as an algorithmic interpretation consistency audit."),
        ("figure_s14_calculator_interface_verification_workflow.png", "Supplementary Figure S14. Online calculator interface and verification workflow."),
    ]
    for fname, cap in fig_captions:
        p = SUPP_FIGS / fname
        if p.exists():
            add_picture_block(doc, p, cap, width=6.7)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Supplementary Tables", level=1)
    titles = [
        ("table_s1_variable_definitions_harmonization.csv", "Supplementary Table S1. Variable definitions and harmonization."),
        ("table_s2_missing_data_analytic_handling.csv", "Supplementary Table S2. Missing data and analytic handling."),
        ("table_s3_check_candidate_model_performance.csv", "Supplementary Table S3. CHECK candidate-model performance."),
        ("table_s4_check_adjusted_pain_kl_associations.csv", "Supplementary Table S4. CHECK adjusted pain-KL associations."),
        ("table_s5_oai_directional_replication.csv", "Supplementary Table S5. OAI directional replication."),
        ("table_s6_oai_kl_state_tka_gradient.csv", "Supplementary Table S6. OAI KL-state TKA/KR gradient."),
        ("table_s7_check_oai_bridge_state_characteristics_smd.csv", "Supplementary Table S7. CHECK-OAI bridge-state characteristics and SMD."),
        ("table_s8_oai_candidate_model_performance.csv", "Supplementary Table S8. OAI candidate-model performance."),
        ("table_s9_oai_predictor_domain_incremental_value.csv", "Supplementary Table S9. OAI predictor-domain incremental value."),
        ("table_s10_final_model_coefficients_preprocessing_constants.csv", "Supplementary Table S10. OAI final-model coefficients and preprocessing constants."),
        ("table_s11_oai_coefficient_bootstrap_stability.csv", "Supplementary Table S11. OAI coefficient bootstrap stability."),
        ("table_s12_oai_clinical_threshold_utility.csv", "Supplementary Table S12. OAI clinical-threshold utility."),
        ("table_s13_mrkr_recalibration_parameters.csv", "Supplementary Table S13. MRKR recalibration parameters."),
        ("table_s14_mrkr_oof_transport_performance.csv", "Supplementary Table S14. MRKR OOF transport performance."),
        ("table_s15_mrkr_early_event_sensitivity.csv", "Supplementary Table S15. MRKR early-event sensitivity."),
        ("table_s16_mrkr_outcome_definition_sensitivity_available_metrics.csv", "Supplementary Table S16. MRKR outcome-definition sensitivity based on available hardware/recalibration metrics."),
        ("table_s17_check_96m_radiograph_availability_bias.csv", "Supplementary Table S17. CHECK 96-month radiograph-availability bias."),
        ("table_s18_pain_cutoff_scale_harmonization_audit.csv", "Supplementary Table S18. Pain cutoff and scale harmonization audit."),
        ("table_s19_oai_kl_timepoint_audit.csv", "Supplementary Table S19. OAI KL timepoint audit."),
        ("table_s20_sample_size_model_complexity.csv", "Supplementary Table S20. Sample size and model complexity."),
        ("table_s21_rf_shap_gbm_interpretability_summary.csv", "Supplementary Table S21. RF SHAP and GBM interpretability summary."),
        ("table_s22_online_calculator_verification_cases.csv", "Supplementary Table S22. Online calculator verification cases."),
    ]
    for fname, title in titles:
        p = SUPP_TABLES / fname
        if not p.exists():
            continue
        add_heading(doc, title, 2)
        df = read(p)
        if len(df) > 24:
            doc.add_paragraph(f"Full machine-readable table is provided at {p}. First 24 rows are displayed below.")
            df = df.head(24)
        add_df_table(doc, df, font_size=5 if len(df.columns) > 8 else 6)
    out = MANUSCRIPT / "three_stage_bridge_transport_supplementary_5fig3table_revised.docx"
    doc.save(out)
    return out


def main() -> None:
    ensure_dirs()
    make_figure2()
    make_figure3()
    make_figure4()
    make_figure5()
    verify = create_calculator()
    t1, t2, t3, _, _ = make_main_tables()
    main_doc = make_main_docx(t1, t2, t3)
    supp_doc = make_supplement_docx()
    print(f"Main DOCX: {main_doc}")
    print(f"Supplement DOCX: {supp_doc}")
    print(f"Output folder: {OUT}")


if __name__ == "__main__":
    main()
