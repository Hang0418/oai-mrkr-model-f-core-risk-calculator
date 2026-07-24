from __future__ import annotations

import csv
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(
    "/Users/hehang/Downloads/时间序列预测模型/膝关节炎/"
    "stage_specific_progression_framework/complete_project_analysis/"
    "three_stage_bridge_transport_manuscript"
)

SOURCE_SUPP = (
    ROOT
    / "JOSR_clinical_translation_revision_20260722/manuscript/"
    "three_stage_bridge_transport_supplementary_materials_JOSR_reordered_preserve_images.docx"
)
SOURCE_MAIN = (
    ROOT
    / "JOSR_clinical_translation_revision_20260722/manuscript/"
    "three_stage_bridge_transport_manuscript_JOSR_clinical_translation_preserve_images.docx"
)
EXTRA_BASE = ROOT / "main_supplement_5fig3table_revision_from_bridge_sample_fixed_base_20260722_v2"
CALCULATOR_FIG = EXTRA_BASE / "supplementary_figures/figure_s14_calculator_interface_verification_workflow.png"
CALCULATOR_TABLE = EXTRA_BASE / "supplementary_tables/table_s22_online_calculator_verification_cases.csv"

OUT_DIR = ROOT / "JOSR_supplement_interleaved_S1_S14_T1_T25_20260722"
MANUSCRIPT_DIR = OUT_DIR / "manuscript"
FIG_DIR = OUT_DIR / "supplementary_figures"
TABLE_DIR = OUT_DIR / "supplementary_tables"
SUPP_OUT = (
    MANUSCRIPT_DIR
    / "three_stage_bridge_transport_supplementary_materials_interleaved_FigureS1-S14_TableS1-S24.docx"
)
MAIN_OUT = (
    MANUSCRIPT_DIR
    / "three_stage_bridge_transport_manuscript_JOSR_clinical_translation_preserve_images_unchanged.docx"
)
MANIFEST = OUT_DIR / "interleaved_supplement_manifest.md"


def text_of(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def set_el_text(el, text: str) -> None:
    text_nodes = list(el.iter(qn("w:t")))
    if not text_nodes:
        return
    text_nodes[0].text = text
    for node in text_nodes[1:]:
        node.text = ""


def is_block_start(text: str) -> bool:
    starts = (
        "Supplementary Figure S",
        "Added Supplementary Figure S",
        "Supplementary Table S",
        "Added Supplementary Table S",
        "Additional Supplementary Materials",
        "JOSR Supplementary Ordering Guide",
        "Supplementary Figures",
        "Supplementary Tables",
        "Supplementary Figure submission order",
        "Supplementary Table submission order",
    )
    return text.startswith(starts)


def collect_blocks(doc: Document) -> dict[str, list]:
    body = list(doc.element.body)
    blocks: dict[str, list] = {}
    idx = 0
    while idx < len(body):
        text = text_of(body[idx])
        if is_block_start(text) and (
            text.startswith("Supplementary Figure S")
            or text.startswith("Added Supplementary Figure S")
            or text.startswith("Supplementary Table S")
            or text.startswith("Added Supplementary Table S")
        ):
            key = text.split(".", 1)[0]
            end = idx + 1
            while end < len(body):
                next_text = text_of(body[end])
                if is_block_start(next_text):
                    break
                end += 1
            blocks[key] = [deepcopy(el) for el in body[idx:end]]
            idx = end
        else:
            idx += 1
    return blocks


def clear_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def append_block(doc: Document, block: list, new_title: str) -> None:
    body = doc.element.body
    cloned = [deepcopy(el) for el in block]
    if cloned:
        set_el_text(cloned[0], new_title)
    sect_pr = body.find(qn("w:sectPr"))
    for el in cloned:
        if sect_pr is not None:
            body.insert(list(body).index(sect_pr), el)
        else:
            body.append(el)


def add_part(doc: Document, title: str) -> None:
    doc.add_paragraph(title, style="Heading 1")


def add_table_from_rows(doc: Document, title: str, headers: list[str], rows: list[list[str]], note: str | None = None) -> None:
    doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        for paragraph in table.rows[0].cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    if note:
        doc.add_paragraph(note)
    doc.add_paragraph("")


def csv_rows(path: Path) -> tuple[list[str], list[list[str]]]:
    with path.open(newline="", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        rows = list(reader)
    return rows[0], rows[1:]


def add_clean_coefficient_stability(doc: Document) -> None:
    path = ROOT / "preserve_existing_images_latest_framework_additions_20260722/added_tables/table_s17_oai_final_model_coefficient_stability.csv"
    headers, rows = csv_rows(path)
    keep = [
        "term",
        "median_coefficient",
        "ci_2_5",
        "ci_97_5",
        "sign_consistency_percent",
        "bootstrap_repetitions",
        "note",
    ]
    idx = [headers.index(col) for col in keep]
    pretty = [
        "Term",
        "Median coefficient",
        "2.5th percentile",
        "97.5th percentile",
        "Sign consistency, %",
        "Bootstrap repetitions",
        "Note",
    ]
    add_table_from_rows(
        doc,
        "Table S16. Bootstrap stability of the OAI final-model coefficients.",
        pretty,
        [[row[i] for i in idx] for row in rows],
        "The ridge model does not perform variable selection; therefore nonzero selection frequency is not reported.",
    )


def add_calculator_cases(doc: Document) -> None:
    headers, rows = csv_rows(CALCULATOR_TABLE)
    pretty = [
        "Calculator module",
        "Verification case",
        "Input values",
        "Source formula probability",
        "Calculator probability",
        "Absolute difference",
    ]
    add_table_from_rows(
        doc,
        "Table S25. Web-calculator verification cases.",
        pretty,
        rows,
        "All displayed test cases showed exact agreement between the source formula and calculator output at the reported precision. Full verification rows can also be supplied as a machine-readable CSV file.",
    )


def add_bridge_state_table(doc: Document) -> None:
    add_table_from_rows(
        doc,
        "Table S10. CHECK 60-month bridge-state distribution.",
        ["Bridge state", "n", "%"],
        [["KL2", "374", "93.5"], ["KL3", "18", "4.5"], ["KL5/TKA", "8", "2.0"]],
        "This bridge-state set was defined separately from the CHECK 24-month prediction endpoint.",
    )


def add_calculator_figure(doc: Document) -> None:
    doc.add_paragraph("Figure S14. Web-based calculator interface and verification workflow.", style="Heading 2")
    doc.add_picture(str(CALCULATOR_FIG), width=Inches(6.5))
    doc.add_paragraph(
        "The research calculator implements the early structural-transition and established arthroplasty-risk modules with fixed verification cases. It is intended for research use, risk communication, and reproducibility checks rather than treatment or surgical selection."
    )
    doc.add_paragraph("")


def style_numbered_titles(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Figure S") or text.startswith("Table S"):
            paragraph.style = "Heading 2"


def doc_stats(path: Path) -> dict[str, int]:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": sum(1 for rel in doc.part._rels.values() if "image" in rel.reltype),
        "figure_titles": sum(1 for p in doc.paragraphs if p.text.strip().startswith("Figure S")),
        "table_titles": sum(1 for p in doc.paragraphs if p.text.strip().startswith("Table S")),
        "added": text.count("Added Supplementary"),
        "gbm": text.count("GBM"),
        "mbkr": text.count("MBKR"),
        "todo": text.lower().count("todo"),
        "tbd": text.lower().count("tbd"),
        "placeholder": text.lower().count("placeholder"),
    }


def build_supplement() -> None:
    MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_MAIN, MAIN_OUT)
    shutil.copy2(CALCULATOR_FIG, FIG_DIR / CALCULATOR_FIG.name)
    shutil.copy2(CALCULATOR_TABLE, TABLE_DIR / CALCULATOR_TABLE.name)

    doc = Document(SOURCE_SUPP)
    blocks = collect_blocks(doc)
    clear_body(doc)

    doc.add_paragraph("Additional file 1: Supplementary figures and tables", style="Title")
    doc.add_paragraph(
        "Supplementary figures and tables are arranged by the manuscript story line. Numbering is unified as Figure S1-Figure S14 and Table S1-Table S24; earlier extra-material labels have been removed."
    )

    add_part(doc, "Part I. Cohort construction and methodological definitions")
    append_block(doc, blocks["Supplementary Figure S1"], "Figure S1. Three-cohort inclusion flow and endpoint-specific analytic populations.")
    append_block(doc, blocks["Supplementary Table S1"], "Table S1. Variable definitions and scale harmonization.")
    append_block(doc, blocks["Supplementary Table S2"], "Table S2. Missing data and analytic handling.")
    append_block(doc, blocks["Supplementary Table S14"], "Table S3. Sample size and model complexity.")
    append_block(doc, blocks["Added Supplementary Table S21"], "Table S4. Pain cutoff and scale-harmonization audit.")
    append_block(doc, blocks["Added Supplementary Table S22"], "Table S5. OAI KL timepoint audit.")

    add_part(doc, "Part II. Early structural-transition prediction in CHECK")
    append_block(doc, blocks["Supplementary Figure S2"], "Figure S2. CHECK KL transition matrices.")
    append_block(doc, blocks["Supplementary Table S3"], "Table S6. CHECK candidate-model performance for incident KL >=2.")
    append_block(doc, blocks["Supplementary Figure S3"], "Figure S3. Calibration of the selected CHECK models across prediction horizons.")
    append_block(doc, blocks["Supplementary Figure S5"], "Figure S4. CHECK candidate-model performance across prediction horizons.")
    append_block(doc, blocks["Supplementary Table S4"], "Table S7. Adjusted pain-KL associations in CHECK.")
    append_block(doc, blocks["Supplementary Table S12"], "Table S8. CHECK 96-month outcome sensitivity analyses.")
    append_block(doc, blocks["Added Supplementary Figure S16"], "Figure S5. CHECK 96-month radiograph-availability bias.")
    append_block(doc, blocks["Added Supplementary Table S20"], "Table S9. Baseline characteristics according to 96-month radiograph availability.")

    add_part(doc, "Part III. CHECK-OAI radiographic bridge analyses")
    add_bridge_state_table(doc)
    append_block(doc, blocks["Supplementary Table S7"], "Table S11. CHECK-OAI characteristics at the radiographic bridge state.")
    append_block(doc, blocks["Supplementary Table S5"], "Table S12. OAI directional replication in baseline KL0/1 knees.")
    append_block(doc, blocks["Added Supplementary Figure S18"], "Figure S6. OAI KL-timepoint sensitivity for the downstream arthroplasty-risk gradient.")
    append_block(doc, blocks["Supplementary Table S6"], "Table S13. OAI KL-state gradient for 60-month TKA/KR risk.")

    add_part(doc, "Part IV. OAI arthroplasty-risk model development")
    append_block(doc, blocks["Supplementary Table S8"], "Table S14. OAI candidate-model performance.")
    append_block(doc, blocks["Added Supplementary Figure S15"], "Figure S7. OAI predictor-domain incremental value.")
    append_block(doc, blocks["Added Supplementary Table S15"], "Table S15. OAI predictor-domain incremental value.")
    append_block(doc, blocks["Added Supplementary Figure S19"], "Figure S8. Bootstrap stability of the OAI final-model coefficients.")
    add_clean_coefficient_stability(doc)
    append_block(doc, blocks["Supplementary Table S13"], "Table S17. Final model equations, coefficients, and preprocessing constants.")
    append_block(doc, blocks["Supplementary Table S10"], "Table S18. OAI predicted-risk tertiles.")
    append_block(doc, blocks["Added Supplementary Table S16"], "Table S19. OAI illustrative clinical-threshold utility.")
    append_block(doc, blocks["Added Supplementary Figure S22"], "Figure S9. Secondary random-forest SHAP interpretation.")

    add_part(doc, "Part V. Transport and recalibration in MRKR")
    append_block(doc, blocks["Supplementary Figure S7"], "Figure S10. Internal target-cohort benchmark performance in MRKR.")
    append_block(doc, blocks["Supplementary Table S9"], "Table S20. MRKR internal target-cohort benchmark performance.")
    append_block(doc, blocks["Added Supplementary Figure S21"], "Figure S11. Patient-grouped out-of-fold MRKR recalibration workflow.")
    append_block(doc, blocks["Supplementary Table S11A"], "Table S21. Fitted MRKR recalibration parameters.")
    append_block(doc, blocks["Supplementary Table S11B"], "Table S22. Out-of-fold MRKR transport and recalibration performance.")
    append_block(doc, blocks["Added Supplementary Figure S17"], "Figure S12. MRKR early-event exclusion sensitivity.")
    append_block(doc, blocks["Added Supplementary Table S24"], "Table S23. MRKR early-event exclusion sensitivity.")
    append_block(doc, blocks["Supplementary Figure S8"], "Figure S13. Exploratory MRKR decision-curve analysis after recalibration.")
    append_block(doc, blocks["Added Supplementary Table S25"], "Table S24. Secondary random-forest SHAP importance.")

    add_part(doc, "Part VI. Research calculator and reproducibility")
    add_calculator_figure(doc)
    add_calculator_cases(doc)

    style_numbered_titles(doc)
    doc.save(SUPP_OUT)

    stats = doc_stats(SUPP_OUT)
    MANIFEST.write_text(
        "\n".join(
            [
                "# Interleaved supplementary figure/table revision",
                "",
                f"Source supplement: `{SOURCE_SUPP}`",
                f"Output supplement: `{SUPP_OUT}`",
                f"Unchanged main manuscript copy: `{MAIN_OUT}`",
                "",
                "Applied ordering:",
                "- Interleaved figures and tables by manuscript story line.",
                "- Used only Figure S1-Figure S14 and Table S1-Table S24 numbering.",
                "- Removed earlier extra-material labels.",
                "- Excluded GBM permutation-importance figure/table and secondary GBM interpretability table from the final ordered supplement.",
                "- Added calculator workflow figure and calculator verification cases from the existing project outputs.",
                "- Did not create an empty MRKR outcome-definition sensitivity table because no completed analysis file was found; Table S24 is therefore assigned to the retained RF SHAP secondary-interpretability table.",
                "",
                f"Verification stats: {stats}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(f"Supplement DOCX: {SUPP_OUT}")
    print(f"Main copy: {MAIN_OUT}")
    print(f"Manifest: {MANIFEST}")
    print(f"Stats: {stats}")


if __name__ == "__main__":
    build_supplement()
