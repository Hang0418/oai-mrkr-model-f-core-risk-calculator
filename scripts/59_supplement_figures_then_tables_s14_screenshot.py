from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(
    "/Users/hehang/Downloads/时间序列预测模型/膝关节炎/"
    "stage_specific_progression_framework/complete_project_analysis/"
    "three_stage_bridge_transport_manuscript"
)

SOURCE_SUPP = (
    ROOT
    / "JOSR_supplement_interleaved_S1_S14_T1_T25_20260722/manuscript/"
    "three_stage_bridge_transport_supplementary_materials_interleaved_FigureS1-S14_TableS1-S24.docx"
)
SOURCE_MAIN = (
    ROOT
    / "JOSR_supplement_interleaved_S1_S14_T1_T25_20260722/manuscript/"
    "three_stage_bridge_transport_manuscript_JOSR_clinical_translation_preserve_images_unchanged.docx"
)
S14_SCREENSHOT = (
    ROOT
    / "enhanced_research_calculator_20260722/screenshots/"
    "figure_s14_enhanced_web_calculator_composite.png"
)

OUT_DIR = ROOT / "JOSR_supplement_figures_then_tables_S1_S14_T1_T25_20260722"
MANUSCRIPT_DIR = OUT_DIR / "manuscript"
FIG_DIR = OUT_DIR / "supplementary_figures"
SUPP_OUT = (
    MANUSCRIPT_DIR
    / "three_stage_bridge_transport_supplementary_materials_FiguresS1-S14_then_TablesS1-S24_S14_screenshot.docx"
)
MAIN_OUT = (
    MANUSCRIPT_DIR
    / "three_stage_bridge_transport_manuscript_JOSR_clinical_translation_preserve_images_unchanged.docx"
)
MANIFEST = OUT_DIR / "figures_then_tables_manifest.md"


def text_of(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def set_el_text(el, text: str) -> None:
    nodes = list(el.iter(qn("w:t")))
    if not nodes:
        return
    nodes[0].text = text
    for node in nodes[1:]:
        node.text = ""


def clear_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def collect_numbered_blocks(doc: Document) -> dict[str, list]:
    body = list(doc.element.body)
    keys = []
    for idx, el in enumerate(body):
        text = text_of(el)
        if text.startswith("Figure S") or text.startswith("Table S"):
            keys.append((idx, text.split(".", 1)[0]))
    blocks: dict[str, list] = {}
    for n, (idx, key) in enumerate(keys):
        end = keys[n + 1][0] if n + 1 < len(keys) else len(body)
        blocks[key] = [deepcopy(el) for el in body[idx:end]]
    return blocks


def append_block(doc: Document, block: list, title: str | None = None) -> None:
    body = doc.element.body
    cloned = []
    for el in block:
        text = text_of(el)
        if text.startswith("Part "):
            continue
        if text in {"Supplementary Figures", "Supplementary Tables"}:
            continue
        cloned.append(deepcopy(el))
    if title and cloned:
        set_el_text(cloned[0], title)
    sect_pr = body.find(qn("w:sectPr"))
    for el in cloned:
        if sect_pr is None:
            body.append(el)
        else:
            body.insert(list(body).index(sect_pr), el)


def add_s14_screenshot(doc: Document) -> None:
    doc.add_paragraph("Figure S14. Web-based calculator interface and verification workflow.", style="Heading 2")
    doc.add_picture(str(S14_SCREENSHOT), width=Inches(6.7))
    doc.add_paragraph(
        "Screenshot of the web-based research calculator opened from the project HTML file. The interface displays the research-use disclaimer, stage-specific model tabs, editable model inputs, and the calculation action. The calculator is intended for research verification and risk communication, not for treatment or surgical selection."
    )
    doc.add_paragraph("")


def doc_stats(path: Path) -> dict[str, int]:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": sum(1 for rel in doc.part._rels.values() if "image" in rel.reltype),
        "figure_titles": sum(1 for p in doc.paragraphs if p.text.strip().startswith("Figure S")),
        "table_titles": sum(1 for p in doc.paragraphs if p.text.strip().startswith("Table S")),
        "part": text.count("Part "),
        "added": text.count("Added Supplementary"),
        "gbm": text.count("GBM"),
        "mbkr": text.count("MBKR"),
        "todo": text.lower().count("todo"),
        "tbd": text.lower().count("tbd"),
        "placeholder": text.lower().count("placeholder"),
    }


def main() -> None:
    MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_MAIN, MAIN_OUT)
    shutil.copy2(S14_SCREENSHOT, FIG_DIR / S14_SCREENSHOT.name)

    source_doc = Document(SOURCE_SUPP)
    blocks = collect_numbered_blocks(source_doc)

    out_doc = Document(SOURCE_SUPP)
    clear_body(out_doc)
    out_doc.add_paragraph("Additional file 1: Supplementary figures and tables", style="Title")
    out_doc.add_paragraph(
        "Supplementary materials are arranged with all figures first, followed by all tables. Numbering is unified as Figure S1-Figure S14 and Table S1-Table S24."
    )

    out_doc.add_paragraph("Supplementary Figures", style="Heading 1")
    for i in range(1, 14):
        append_block(out_doc, blocks[f"Figure S{i}"])
    add_s14_screenshot(out_doc)

    out_doc.add_paragraph("Supplementary Tables", style="Heading 1")
    for i in range(1, 26):
        append_block(out_doc, blocks[f"Table S{i}"])

    out_doc.save(SUPP_OUT)
    stats = doc_stats(SUPP_OUT)
    MANIFEST.write_text(
        "\n".join(
            [
                "# Supplement figures-then-tables revision",
                "",
                f"Source supplement: `{SOURCE_SUPP}`",
                f"Output supplement: `{SUPP_OUT}`",
                f"Unchanged main manuscript copy: `{MAIN_OUT}`",
                f"Figure S14 screenshot: `{S14_SCREENSHOT}`",
                "",
                "Applied changes:",
                "- Removed Part sections.",
                "- Placed Figure S1-Figure S14 together before Table S1-Table S24.",
                "- Replaced Figure S14 with a screenshot captured from the project web-calculator HTML.",
                "- Preserved the existing Figure S1-S13 and Table S1-S25 content from the prior unified-numbering supplement.",
                "",
                f"Verification stats: {stats}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(f"Supplement DOCX: {SUPP_OUT}")
    print(f"Main copy: {MAIN_OUT}")
    print(f"S14 screenshot: {S14_SCREENSHOT}")
    print(f"Manifest: {MANIFEST}")
    print(f"Stats: {stats}")


if __name__ == "__main__":
    main()
