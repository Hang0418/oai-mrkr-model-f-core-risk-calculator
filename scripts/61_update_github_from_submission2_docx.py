from __future__ import annotations

import csv
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from PIL import Image


ROOT = Path("/Users/hehang/Downloads/时间序列预测模型/膝关节炎")
SUBMISSION_DIR = ROOT / "投稿2"
MAIN_DOCX = SUBMISSION_DIR / "Main Tables and Figure.docx"
SUPP_DOCX = SUBMISSION_DIR / "supplementary_materials.docx"

REPO = (
    ROOT
    / "public_release/oai-mrkr-model-f-core-risk-calculator"
)

MAIN_FIG_NAMES = [
    "figure1_stage_specific_framework",
    "figure2_check_early_structural_transition",
    "figure3_radiographic_bridge_arthroplasty_gradient",
    "figure4_oai_arthroplasty_model_stratification",
    "figure5_mrkr_transport_recalibration_sensitivity",
]

# Submission 2 has 14 supplementary figure captions but 13 embedded image files.
# Figure S9 (secondary RF SHAP) is retained from the prior repository version.
SUPP_IMAGE_TO_NAME = {
    1: "figure_s1_three_cohort_inclusion_flow",
    2: "figure_s2_check_kl_transition_matrices",
    3: "figure_s3_check_calibration",
    4: "figure_s4_check_candidate_model_performance",
    5: "figure_s5_check_96m_radiograph_availability_bias",
    6: "figure_s6_oai_kl_timepoint_sensitivity",
    7: "figure_s7_oai_predictor_domain_incremental_value",
    8: "figure_s8_oai_coefficient_stability",
    9: "figure_s10_mrkr_internal_benchmark",
    10: "figure_s11_mrkr_oof_recalibration_workflow",
    11: "figure_s12_mrkr_early_event_exclusion_sensitivity",
    12: "figure_s13_mrkr_decision_curve",
    13: "figure_s14_web_calculator_composite",
}


def media_images(docx: Path) -> list[bytes]:
    with ZipFile(docx) as zf:
        names = [
            n
            for n in zf.namelist()
            if n.startswith("word/media/image") and n.lower().endswith(".png")
        ]
        names = sorted(
            names,
            key=lambda n: int("".join(filter(str.isdigit, Path(n).stem)) or "0"),
        )
        return [zf.read(n) for n in names]


def write_png_and_pdf(data: bytes, stem: Path) -> None:
    stem.parent.mkdir(parents=True, exist_ok=True)
    png = stem.with_suffix(".png")
    png.write_bytes(data)
    pdf = stem.with_suffix(".pdf")
    try:
        image = Image.open(png).convert("RGB")
        image.save(pdf, "PDF", resolution=300.0)
    except Exception:
        if pdf.exists():
            pdf.unlink()


def table_rows(table) -> list[list[str]]:
    return [[cell.text.replace("\n", " ").strip() for cell in row.cells] for row in table.rows]


def write_csv(path: Path, rows: list[list[str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        csv.writer(f).writerows(rows)


def update_tables(docx: Path, target: Path, prefix: str) -> int:
    doc = Document(docx)
    for existing in target.glob(f"{prefix}*.csv"):
        existing.unlink()
    for idx, table in enumerate(doc.tables, 1):
        write_csv(target / f"{prefix}{idx:02d}.csv", table_rows(table))
    return len(doc.tables)


def apply_final_consistency_revision() -> int:
    """Apply post-extraction table fixes from the final submission audit."""
    target = REPO / "tables/supplementary"

    table_s3 = target / "table_s03.csv"
    rows = list(csv.reader(table_s3.open(encoding="utf-8-sig")))
    for row in rows[1:]:
        if row and row[0] == "CHECK":
            row[5] = "8 predictor parameters plus intercept"
        elif row and row[0] == "OAI":
            row[5] = "7 predictor parameters plus intercept"
    write_csv(table_s3, rows)

    table_s11 = target / "table_s11.csv"
    rows = list(csv.reader(table_s11.open(encoding="utf-8-sig")))
    if rows and len(rows[0]) > 1:
        rows[0][1] = "CHECK 60-month bridge-state set, n=400"
    write_csv(table_s11, rows)

    table_s24 = target / "table_s24.csv"
    table_s25 = target / "table_s25.csv"
    if table_s25.exists():
        table_s24.unlink(missing_ok=True)
        table_s25.rename(table_s24)

    note = target / "README_table_numbering_note.md"
    note.write_text(
        "Supplementary CSVs follow the final consistency revision: the prior "
        "Table S24 was removed, and web-calculator verification was renumbered "
        "as Table S24. No Table S25 is retained in the public CSV release.\n",
        encoding="utf-8",
    )

    return len(list(target.glob("table_s*.csv")))


def update_figures() -> tuple[int, int]:
    main_images = media_images(MAIN_DOCX)
    if len(main_images) != 5:
        raise ValueError(f"Expected 5 main figures, found {len(main_images)}")
    for data, name in zip(main_images, MAIN_FIG_NAMES):
        write_png_and_pdf(data, REPO / "figures/main" / name)

    supp_images = media_images(SUPP_DOCX)
    if len(supp_images) != 13:
        raise ValueError(f"Expected 13 embedded supplementary images, found {len(supp_images)}")
    for idx, data in enumerate(supp_images, 1):
        write_png_and_pdf(data, REPO / "figures/supplementary" / SUPP_IMAGE_TO_NAME[idx])
    return len(main_images), len(supp_images)


def update_readme_note() -> None:
    readme = REPO / "README.md"
    text = readme.read_text(encoding="utf-8")
    note = (
        "\n## Submission asset source\n\n"
        "The current figure and table assets were refreshed from the `投稿2` Word files: "
        "`Main Tables and Figure.docx` and `supplementary_materials.docx`. "
        "Only extracted figure images and CSV table summaries are included in this repository; "
        "the Word files themselves and raw cohort data are not uploaded.\n"
    )
    marker = "\n## Submission asset source\n"
    if marker in text:
        text = text.split(marker)[0].rstrip() + note
    else:
        text = text.rstrip() + note
    readme.write_text(text, encoding="utf-8")


def main() -> None:
    main_count, supp_count = update_figures()
    main_tables = update_tables(MAIN_DOCX, REPO / "tables/main", "table")
    supp_tables = update_tables(SUPP_DOCX, REPO / "tables/supplementary", "table_s")
    retained_supp_tables = apply_final_consistency_revision()
    update_readme_note()
    manifest = REPO / "MANIFEST.md"
    manifest.write_text(
        "\n".join(
            [
                "# Release manifest",
                "",
                "Included assets are restricted to manuscript figure/table outputs and calculator code.",
                "",
                "Latest source: submission 2 Word files.",
                f"- Main figures extracted from `{MAIN_DOCX}`: {main_count}",
                f"- Supplementary embedded figures extracted from `{SUPP_DOCX}`: {supp_count}",
                "- Figure S9 is retained from the existing repository because the submission 2 supplement has a Figure S9 caption but no separate embedded image for that panel.",
                f"- Main tables extracted from `{MAIN_DOCX}`: {main_tables}",
                f"- Supplementary tables extracted from `{SUPP_DOCX}`: {supp_tables} before consistency revision; {retained_supp_tables} retained in the public CSV release",
                "- Table S3, Table S11, and final web-calculator table numbering were harmonized after extraction according to the final submission consistency revision.",
                "",
                "No raw cohort data, DOCX manuscript files, WPS backups, or intermediate patient-level source files are included.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(f"main_figures={main_count} supplementary_embedded_figures={supp_count}")
    print(f"main_tables={main_tables} supplementary_tables={retained_supp_tables}")


if __name__ == "__main__":
    main()
